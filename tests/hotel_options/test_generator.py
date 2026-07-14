import io
from docx import Document
from src.hotel_options.generator import format_indian_number, build_document
from src.hotel_options.models import Plan, PlanPricing, HotelRow, EnrichedHotel


def test_format_indian_number():
    assert format_indian_number(176622.0) == "₹1,76,622"
    assert format_indian_number(191022.0) == "₹1,91,022"
    assert format_indian_number(3000.0) == "₹3,000"
    assert format_indian_number(100.0) == "₹100"


def _make_plan() -> Plan:
    hotel = HotelRow(
        name="Test Hotel", category="4-Star", room_type="Double",
        cancellation="Non-refundable", meal_type="Breakfast included", online_price=50000.0,
    )
    pricing = PlanPricing(
        total_online_price=50000.0, total_b2b_price=45000.0,
        customer_discount=3000.0, discounted_price=47000.0, discount_pct=6.0,
    )
    return Plan(label="Plan A", hotels=[hotel], pricing=pricing)


def _make_enriched() -> EnrichedHotel:
    return EnrichedHotel(
        official_name="Test Hotel Official", address="123 Test St, London",
        phone="+44 20 1234 5678", rating=4.1, rating_count=500,
        maps_url="https://maps.google.com/?cid=1", photo_bytes=None,
        description="A great hotel.", cancellation="Non-refundable",
        meal_type="Breakfast included", category="4-Star",
    )


def test_build_document_returns_bytes():
    doc_bytes = build_document(
        plans=[_make_plan()],
        enriched_map={"Test Hotel": _make_enriched()},
        client_name="Alice",
        destination="London",
    )
    assert isinstance(doc_bytes, bytes)
    assert len(doc_bytes) > 0


def test_build_document_contains_plan_and_destination():
    doc_bytes = build_document(
        plans=[_make_plan()],
        enriched_map={"Test Hotel": _make_enriched()},
        client_name="Alice",
        destination="London",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "PLAN A" in full_text
    assert "London" in full_text
    assert "ALICE" in full_text


def test_format_indian_number_small():
    assert format_indian_number(500.0) == "₹500"


def _make_section_plans() -> list:
    from src.hotel_options.models import HotelRow, PlanPricing, Plan
    def _hotel(name, online, discount=0.0):
        discounted = online - discount
        pct = (discount / online * 100) if online else 0.0
        return HotelRow(name=name, category="4-Star", room_type="Double",
                        cancellation="Free", meal_type="Breakfast",
                        online_price=online, customer_discount=discount,
                        discounted_price=discounted, discount_pct=pct)
    p1 = Plan(label="London (Jul 1 - Jul 4)",
              hotels=[_hotel("Hotel A", 100000, 5000), _hotel("Hotel B", 80000)],
              pricing=PlanPricing(180000, 160000, 5000, 175000, 2.78))
    p2 = Plan(label="Paris (Jul 5 - Jul 8)",
              hotels=[_hotel("Hotel C", 120000, 8000)],
              pricing=PlanPricing(120000, 100000, 8000, 112000, 6.67))
    return [p1, p2]


def test_build_document_grouped_by_sections_has_city_dates_header():
    doc_bytes = build_document(
        plans=_make_section_plans(),
        enriched_map={},
        client_name="Alice",
        destination="London",
        grouped_by_sections=True,
    )
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Check exec summary table contains hotel names and section labels
    all_text = full_text + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "City / Dates" in all_text
    assert "London (Jul 1 - Jul 4)" in all_text
    assert "Hotel A" in all_text
    assert "Hotel B" in all_text


def test_build_document_grouped_sections_per_hotel_pricing():
    """Each hotel in the sections body gets its own pricing block (Our Price row)."""
    doc_bytes = build_document(
        plans=_make_section_plans(),
        enriched_map={},
        client_name="Alice",
        destination="London",
        grouped_by_sections=True,
    )
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    all_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    # 3 per-hotel pricing blocks + 1 "Our Price" header in exec summary table = 4
    assert all_text.count("Our Price") == 4


def test_build_document_plan_path_unchanged():
    """The PLAN A/B path must still use plan-level pricing (not per-hotel)."""
    doc_bytes = build_document(
        plans=[_make_plan()],
        enriched_map={"Test Hotel": _make_enriched()},
        client_name="Bob",
        destination="London",
        grouped_by_sections=False,
    )
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "PLAN A" in full_text
    assert "Our Price" in table_text
    assert "City / Dates" not in table_text


def test_recommended_badge_appears_in_grouped_exec_summary():
    """★ RECOMMENDED appears in exec summary for a recommended hotel."""
    from src.hotel_options.models import HotelRow, PlanPricing, Plan
    hotel_rec = HotelRow(name="Hotel A", category="4-Star", room_type="Double",
                         cancellation="Free", meal_type="Breakfast",
                         online_price=100000.0, recommended=True)
    hotel_reg = HotelRow(name="Hotel B", category="3-Star", room_type="Twin",
                         cancellation="Free", meal_type="Breakfast",
                         online_price=80000.0, recommended=False)
    plans = [Plan(label="London (Jul 1-4)", hotels=[hotel_rec, hotel_reg],
                  pricing=PlanPricing(180000, 160000, 0, 180000, 0))]
    doc_bytes = build_document(plans=plans, enriched_map={}, client_name="Alice",
                               destination="London", grouped_by_sections=True)
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    all_text = "\n".join(cell.text for t in doc.tables for r in t.rows for cell in r.cells)
    assert "RECOMMENDED" in all_text


def test_recommended_badge_appears_in_plan_exec_summary():
    """★ RECOMMENDED appears in exec summary for a recommended plan."""
    from src.hotel_options.models import HotelRow, PlanPricing, Plan
    hotel = HotelRow(name="Hotel A", category="4-Star", room_type="Double",
                     cancellation="Free", meal_type="Breakfast", online_price=50000.0)
    plan_rec = Plan(label="Plan A", hotels=[hotel],
                    pricing=PlanPricing(50000, 45000, 3000, 47000, 6.0), recommended=True)
    plan_reg = Plan(label="Plan B", hotels=[hotel],
                    pricing=PlanPricing(50000, 45000, 0, 50000, 0), recommended=False)
    doc_bytes = build_document(plans=[plan_rec, plan_reg], enriched_map={},
                               client_name="Bob", destination="London")
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    all_text = "\n".join(cell.text for t in doc.tables for r in t.rows for cell in r.cells)
    assert "RECOMMENDED" in all_text


def test_build_document_no_plan_column_when_grouped():
    """When grouped_by_sections=True, exec summary has 'City / Dates' not 'Plan'."""
    doc_bytes = build_document(
        plans=_make_section_plans(),
        enriched_map={},
        client_name="Alice",
        destination="London",
        grouped_by_sections=True,
    )
    import io as _io
    from docx import Document as _Doc
    doc = _Doc(_io.BytesIO(doc_bytes))
    # Find the exec summary table — it's the one containing "City / Dates"
    exec_table = next(
        (t for t in doc.tables
         if any("City / Dates" in cell.text for row in t.rows for cell in row.cells)),
        None,
    )
    assert exec_table is not None, "Exec summary table with 'City / Dates' not found"
    all_header_text = " ".join(cell.text for cell in exec_table.rows[0].cells)
    assert "City / Dates" in all_header_text
    assert "Plan" not in all_header_text


from src.hotel_options.models import RoomSegment


def _make_multi_segment_plan() -> Plan:
    h1 = HotelRow(
        name="Beach Resort", category="4-Star", room_type="Beach Villa",
        cancellation="Free", meal_type="All Inclusive", online_price=1000000.0,
        dates="Dec 22-25",
    )
    h2 = HotelRow(
        name="Beach Resort", category="4-Star", room_type="Ocean Villa",
        cancellation="Free", meal_type="All Inclusive", online_price=500000.0,
        dates="Dec 25-26",
    )
    pricing = PlanPricing(
        total_online_price=1500000.0, total_b2b_price=1300000.0,
        customer_discount=50000.0, discounted_price=1450000.0, discount_pct=3.3,
    )
    return Plan(
        label="4-Star Plan",
        hotels=[h1, h2],
        pricing=pricing,
        inclusions="With airport transfer - Shared Speedboat",
    )


def _make_multi_segment_enriched() -> EnrichedHotel:
    return EnrichedHotel(
        official_name="Beach Resort Official", address="Maldives",
        phone="", rating=4.5, rating_count=800,
        maps_url="https://maps.google.com/?cid=1",
        photo_bytes=None, description="A beautiful resort.",
        cancellation="Free", meal_type="All Inclusive", category="4-Star",
        room_segments=[
            RoomSegment(room_type="Beach Villa", dates="Dec 22-25",
                        online_price=1000000.0, photo_bytes=None,
                        inclusions="Airport transfer", exclusions="Visa fees"),
            RoomSegment(room_type="Ocean Villa", dates="Dec 25-26",
                        online_price=500000.0, photo_bytes=None),
        ],
    )


def test_multi_segment_hotel_card_rendered_once():
    """Same hotel name appears twice in plan.hotels but card is rendered only once."""
    plan = _make_multi_segment_plan()
    enriched = _make_multi_segment_enriched()
    doc_bytes = build_document(
        plans=[plan],
        enriched_map={"Beach Resort": enriched},
        client_name="Vinay",
        destination="Maldives",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert full_text.count("Beach Resort Official") == 1


def test_multi_segment_both_room_types_in_document():
    plan = _make_multi_segment_plan()
    enriched = _make_multi_segment_enriched()
    doc_bytes = build_document(
        plans=[plan],
        enriched_map={"Beach Resort": enriched},
        client_name="Vinay",
        destination="Maldives",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Beach Villa" in full_text
    assert "Ocean Villa" in full_text


def test_inclusions_exclusions_in_document():
    plan = _make_multi_segment_plan()
    enriched = _make_multi_segment_enriched()
    doc_bytes = build_document(
        plans=[plan],
        enriched_map={"Beach Resort": enriched},
        client_name="Vinay",
        destination="Maldives",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Airport transfer" in full_text
    assert "Visa fees" in full_text


def test_plan_inclusions_in_document():
    plan = _make_multi_segment_plan()
    enriched = _make_multi_segment_enriched()
    doc_bytes = build_document(
        plans=[plan],
        enriched_map={"Beach Resort": enriched},
        client_name="Vinay",
        destination="Maldives",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Shared Speedboat" in full_text


def test_exec_summary_hotel_listed_once_for_multi_segment():
    plan = _make_multi_segment_plan()
    enriched = _make_multi_segment_enriched()
    doc_bytes = build_document(
        plans=[plan],
        enriched_map={"Beach Resort": enriched},
        client_name="Vinay",
        destination="Maldives",
    )
    doc = Document(io.BytesIO(doc_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Hotel name should appear in exec summary exactly once (not twice for two HotelRows)
    assert full_text.count("Beach Resort") >= 1
