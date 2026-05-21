from docx import Document
from src.aig.styles import ensure_styles, AIG_STYLES


def test_ensure_styles_adds_all_aig_styles():
    doc = Document()
    ensure_styles(doc)
    style_names = {s.name for s in doc.styles}
    for name, *_ in AIG_STYLES:
        assert name in style_names, f"Missing style: {name}"


def test_ensure_styles_idempotent():
    doc = Document()
    ensure_styles(doc)
    ensure_styles(doc)  # second call must not raise
    style_names = {s.name for s in doc.styles}
    aig_names = [name for name, *_ in AIG_STYLES]
    assert len([n for n in aig_names if n in style_names]) == len(AIG_STYLES)


def test_aig_title_is_bold_and_large():
    doc = Document()
    ensure_styles(doc)
    style = doc.styles["AIG Title"]
    assert style.font.bold is True
    assert style.font.size.pt == 28
