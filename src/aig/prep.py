"""AIG prep: generate library context and client profile companion documents."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


# ---------------------------------------------------------------------------
# Country lookup: aig-library folder name → clean country/region label
# ---------------------------------------------------------------------------

FOLDER_TO_COUNTRY: dict[str, str] = {
    "Africa": "Africa",
    "Almaty Kazakhstan": "Kazakhstan",
    "Andamans": "India",
    "Australia": "Australia",
    "Austria": "Austria",
    "Azerbaijan and Georgia": "Azerbaijan / Georgia",
    "Bali Indonesia": "Indonesia",
    "Belgium": "Belgium",
    "Bhutan": "Bhutan",
    "China": "China",
    "Dubai and Abu Dhabi - UAE": "UAE",
    "England and Scotland": "United Kingdom",
    "France": "France",
    "French Polynesia": "French Polynesia",
    "Germany": "Germany",
    "Goa": "India",
    "Greece and Turkey": "Greece / Turkey",
    "Italy": "Italy",
    "Japan": "Japan",
    "Karnataka, Kerala and Tamil Nadu": "India",
    "Maldives": "Maldives",
    "Mauritius": "Mauritius",
    "Netherlands": "Netherlands",
    "North India": "India",
    "Northeast India": "India",
    "Philippines": "Philippines",
    "Qatar": "Qatar",
    "Rajasthan": "India",
    "Singapore and Malaysia": "Singapore / Malaysia",
    "South America": "South America",
    "South Korea": "South Korea",
    "Spain": "Spain",
    "Srilanka": "Sri Lanka",
    "Switzerland": "Switzerland",
    "Thailand": "Thailand",
    "USA": "USA",
    "Vietnam": "Vietnam",
}

# Sections from library JSON to include (order determines output order).
# Attractions and hotels are excluded per spec.
_LIBRARY_SECTIONS = [
    ("restaurants",        "Curated Restaurants"),
    ("local_dishes",       "Must-Try Local Dishes"),
    ("souvenirs",          "Souvenir Shopping"),
    ("transport_options",  "Getting Around"),
    ("safety_tips",        "Safety Tips"),
    ("connectivity_tips",  "Mobile Connectivity"),
    ("emergency_contacts", "Emergency Contacts"),
    ("health_tips",        "Health & Vaccination"),
]


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class PrepContext:
    client_name: str
    destination_label: str    # e.g. "London" or "Lima, Cusco, Sacred Valley"
    cities: list[str]
    date_range: str           # e.g. "28 Jun 2026 – 4 Jul 2026" or ""
    hotels: dict[str, str]    # city → hotel name
    dietary_notes: str        # joined string, e.g. "vegetarian, chicken, seafood"
    transport_mode: str       # e.g. "Public Transport" or ""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _format_iso_date(iso: Optional[str]) -> str:
    """Convert 'YYYY-MM-DD' to '28 Jun 2026', or return '' on failure."""
    if not iso:
        return ""
    try:
        from datetime import date
        d = date.fromisoformat(iso)
        return d.strftime("%-d %b %Y")
    except (ValueError, TypeError):
        return iso  # return as-is if unparseable


def _join_diet(restrictions: list[str], allergies: list[str], preferences: list[str]) -> str:
    combined = [x for x in restrictions + allergies + preferences if x]
    return ", ".join(combined)


def _build_prep_context_from_trip_facts(data: dict) -> PrepContext:
    """Build PrepContext from a trip_facts.json dict (TripFacts schema)."""
    client_names: list[str] = data.get("client_names") or []
    client_name = client_names[0] if client_names else ""

    cities: list[str] = data.get("destinations") or []
    destination_label = ", ".join(cities) if cities else ""

    start = _format_iso_date(data.get("trip_start_date"))
    end = _format_iso_date(data.get("trip_end_date"))
    date_range = f"{start} – {end}" if start and end else start or end

    hotels: dict[str, str] = {}
    for h in data.get("hotels") or []:
        city = h.get("city") or ""
        name = h.get("hotel_name") or ""
        if city and name:
            hotels[city] = name

    transport_mode = data.get("local_transport") or ""

    dietary_notes = _join_diet(
        data.get("dietary_restrictions") or [],
        data.get("food_allergies") or [],
        data.get("cuisine_preferences") or [],
    )

    return PrepContext(
        client_name=client_name,
        destination_label=destination_label,
        cities=cities,
        date_range=date_range,
        hotels=hotels,
        dietary_notes=dietary_notes,
        transport_mode=transport_mode,
    )


def _build_prep_context_from_itinerary_data(itinerary) -> PrepContext:
    """Build PrepContext from an ItineraryData object (single-file parser output)."""
    client_name = getattr(itinerary, "client_name", None) or ""

    cities: list[str] = list(getattr(itinerary, "destinations", None) or [])
    if not cities:
        dest = getattr(itinerary, "destination", None)
        if dest and dest != "Unknown":
            cities = [dest]
    destination_label = ", ".join(cities)

    start = _format_iso_date(getattr(itinerary, "trip_start_date", None))
    end = _format_iso_date(getattr(itinerary, "trip_end_date", None))
    date_range = f"{start} – {end}" if start and end else start or end

    hotels: dict[str, str] = {}
    for stay in getattr(itinerary, "hotel_stays", None) or []:
        city = getattr(stay, "city", None) or ""
        name = getattr(stay, "hotel_name", None) or ""
        if city and name:
            hotels[city] = name
    if not hotels:
        hotel_name = getattr(itinerary, "hotel_name", None)
        if hotel_name and cities:
            hotels[cities[0]] = hotel_name

    transport_mode = getattr(itinerary, "transport_mode", None) or ""

    dietary_notes = _join_diet(
        list(getattr(itinerary, "dietary_preferences", None) or []),
        list(getattr(itinerary, "food_allergies", None) or []),
        list(getattr(itinerary, "cuisine_preferences", None) or []),
    )

    return PrepContext(
        client_name=client_name,
        destination_label=destination_label,
        cities=cities,
        date_range=date_range,
        hotels=hotels,
        dietary_notes=dietary_notes,
        transport_mode=transport_mode,
    )


def _clean_hotel_name(raw: str) -> str:
    """Strip room-type suffixes and formatting noise from a raw table cell."""
    # Remove parenthetical room type, e.g. "(Deluxe Room)"
    name = re.sub(r"\s*\([^)]*\)\s*", " ", raw).strip()
    # Replace newlines (soft line-wraps inside the cell) with spaces
    name = re.sub(r"\s+", " ", name.replace("\n", " ")).strip()
    # Remove digit-based room-type code, e.g. "3D/2N"
    name = re.split(r"\d+[A-Z]/\d+[A-Z]", name)[0].strip()
    # Split at camelCase boundary, e.g. "MirafloresSuperior" → "Miraflores"
    name = re.split(r"(?<=[a-z])(?=[A-Z])", name)[0].strip()
    # Title-case if the raw value was ALL CAPS
    if name == name.upper():
        return name.title()
    return name


def _extract_from_docx(input_path: Path) -> dict:
    """Extract trip context directly from a BVM DOCX file.

    Handles two BVM formats:
    - Service voucher: tables with CITY_NAME/Hotel columns, CLIENT table with
      MR./MRS./MS. prefix, ARRIVAL DATE / DEPARTURE DATE flight tables, and
      a "Guest Name: X" paragraph.
    - Notes file: freeform paragraphs with a destination title on the first
      line, "Day N | DD Month" day headers, and "Arrival in X" / "Overnight: X"
      location cues.

    Returns a dict with whichever keys could be extracted; missing keys are
    absent (not None) so the caller can apply its own defaults.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return {}

    try:
        doc = Document(str(input_path))
    except Exception:
        return {}

    result: dict = {}

    # Collect all paragraph text (each paragraph may contain soft line-breaks
    # via <w:br/>, which python-docx exposes as \n within p.text).
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paras)

    # Collect table data: list of list-of-rows, each row is list of cell strings.
    tables = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in t.rows]
        tables.append(rows)

    # ------------------------------------------------------------------ #
    # 1. Client name                                                       #
    # ------------------------------------------------------------------ #

    # Priority A: "Guest Name: X" somewhere in paragraph text
    m = re.search(r"Guest Name[:\s]*([^\n]+)", full_text, re.IGNORECASE)
    if m:
        result["client_name"] = m.group(1).strip()

    # Priority B: "MR./MRS./MS./DR. FIRSTNAME [LASTNAME]" in a table cell
    if "client_name" not in result:
        title_re = re.compile(
            r"^(?:MR\.|MRS\.|MS\.|DR\.|MASTER)\s+([A-Z][A-Z\s]+)$"
        )
        for rows in tables:
            for row in rows[:3]:
                for cell in row:
                    m2 = title_re.match(cell.strip())
                    if m2:
                        result["client_name"] = m2.group(1).strip().title()
                        break
                if "client_name" in result:
                    break
            if "client_name" in result:
                break

    # Priority C: filename stem — first valid name segment (expanded reject list;
    # handles multi-word first segments by trying the first word).
    if "client_name" not in result:
        stem = input_path.stem
        _reject = {
            "service", "voucher", "all", "inclusive", "guide", "itinerary",
            "final", "updated", "new", "draft", "for", "aig", "notes",
            "input", "client",
        }
        for sep in ("-", "_"):
            parts = stem.split(sep)
            if len(parts) < 2:
                continue
            first_seg = parts[0].strip()
            # Try single-word segment; if multi-word, try only the first word.
            candidates = [first_seg] if " " not in first_seg else first_seg.split()[:1]
            for word in candidates:
                word = word.strip()
                if (word.lower() not in _reject
                        and re.match(r"^[A-Za-z]+$", word)
                        and len(word) >= 3):
                    result["client_name"] = word.capitalize()
                    break
            if "client_name" in result:
                break

    # ------------------------------------------------------------------ #
    # 2. Dates                                                             #
    # ------------------------------------------------------------------ #

    _MONTHS = (
        "January|February|March|April|May|June|July|"
        "August|September|October|November|December|"
        "Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
    )
    _date_re = re.compile(
        r"(\d{1,2})\s+(" + _MONTHS + r")(?:\s+(\d{4}))?",
        re.IGNORECASE,
    )

    def _parse_date(text: str) -> Optional[str]:
        dm = _date_re.search(text)
        if dm:
            day = int(dm.group(1))
            month = dm.group(2)[:3].capitalize()
            year = dm.group(3) or "2026"
            return f"{day} {month} {year}"
        return None

    # From paragraph text: "Date of Travel: DD Month YYYY"
    m3 = re.search(r"Date of Travel[:\s]*([^\n]+)", full_text, re.IGNORECASE)
    if m3:
        d = _parse_date(m3.group(1))
        if d:
            result["trip_start_date"] = d

    m4 = re.search(r"Departure Date[:\s]*([^\n]+)", full_text, re.IGNORECASE)
    if m4:
        d = _parse_date(m4.group(1))
        if d:
            result["trip_end_date"] = d

    # From tables with ARRIVAL DATE / DEPARTURE DATE header rows
    for rows in tables:
        if not rows:
            continue
        header_upper = [c.upper() for c in rows[0]]
        if "ARRIVAL DATE" in header_upper and "trip_start_date" not in result and len(rows) > 1:
            col = header_upper.index("ARRIVAL DATE")
            if col < len(rows[1]):
                d = _parse_date(rows[1][col])
                if d:
                    result["trip_start_date"] = d
        if "DEPARTURE DATE" in header_upper and "trip_end_date" not in result and len(rows) > 1:
            col = header_upper.index("DEPARTURE DATE")
            if col < len(rows[1]):
                d = _parse_date(rows[1][col])
                if d:
                    result["trip_end_date"] = d

    # From notes format: scan all para dates, use first/last as fallback
    if "trip_start_date" not in result:
        all_dates = [_parse_date(p) for p in paras]
        all_dates = [d for d in all_dates if d]
        if all_dates:
            result["trip_start_date"] = all_dates[0]
            if len(all_dates) > 1:
                result["trip_end_date"] = all_dates[-1]

    # ------------------------------------------------------------------ #
    # 3. Cities and hotels from tables                                     #
    # ------------------------------------------------------------------ #

    cities: list[str] = []
    hotels: dict[str, str] = {}

    for rows in tables:
        if not rows:
            continue
        # Look for a header row containing a city column in the first 3 rows.
        city_col = hotel_col = header_row_idx = None
        for i, row in enumerate(rows[:3]):
            norm = [c.upper().replace(" ", "_") for c in row]
            if "CITY_NAME" in norm or "CITY" in norm:
                header_row_idx = i
                city_col = norm.index("CITY_NAME") if "CITY_NAME" in norm else norm.index("CITY")
                for key in ("HOTEL_NAME", "HOTEL"):
                    if key in norm:
                        hotel_col = norm.index(key)
                        break
                break

        if header_row_idx is None or city_col is None:
            continue

        for row in rows[header_row_idx + 1:]:
            if city_col >= len(row):
                continue
            city_raw = row[city_col].strip()
            if not city_raw or city_raw.upper() in ("CITY_NAME", "CITY"):
                continue
            # Strip parenthetical qualifiers, e.g. "Lima (2nd Entrance)"
            city = re.sub(r"\s*\([^)]*\)\s*$", "", city_raw).strip().title()
            if city and city not in cities:
                cities.append(city)
            if hotel_col is not None and hotel_col < len(row) and city not in hotels:
                hotel_raw = row[hotel_col].strip()
                hotel = _clean_hotel_name(hotel_raw)
                if hotel and hotel.upper() not in ("HOTEL_NAME", "HOTEL"):
                    hotels[city] = hotel

    if cities:
        result["cities"] = cities
    if hotels:
        result["hotels"] = hotels

    # ------------------------------------------------------------------ #
    # 4. Cities from notes-format paragraphs (if table extraction failed) #
    # ------------------------------------------------------------------ #

    if "cities" not in result:
        notes_cities: list[str] = []

        # Title line: "Dehradun & Mussoorie Family Getaway" → [Dehradun, Mussoorie]
        if paras:
            title = paras[0]
            title_clean = re.sub(
                r"\s+(?:Family|Getaway|Adventure|Tour|Trip|Holiday|Vacation|"
                r"Package|Programme|Escape|Journey|Experience)\b.*$",
                "", title, flags=re.IGNORECASE,
            ).strip()
            for part in re.split(r"\s+(?:&|and)\s+", title_clean, flags=re.IGNORECASE):
                part = part.strip()
                if part and len(part) >= 3 and re.match(r"^[A-Za-z\s]+$", part):
                    notes_cities.append(part)

        # "Arrival in X" — often in the first day header
        for para in paras[:15]:
            m5 = re.search(r"Arrival in ([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)", para)
            if m5:
                city = m5.group(1).strip()
                if city not in notes_cities:
                    notes_cities.append(city)

        if notes_cities:
            result["cities"] = notes_cities

    return result


def _build_prep_context_from_extracted(extracted: dict) -> PrepContext:
    """Build PrepContext from the dict returned by _extract_from_docx."""
    client_name = extracted.get("client_name") or ""
    cities: list[str] = list(extracted.get("cities") or [])
    destination_label = ", ".join(cities)

    # Dates from the extractor are already human-readable ("13 Aug 2026")
    start = extracted.get("trip_start_date") or ""
    end = extracted.get("trip_end_date") or ""
    date_range = f"{start} – {end}" if start and end else start or end

    hotels: dict[str, str] = dict(extracted.get("hotels") or {})
    transport_mode = extracted.get("transport_mode") or ""
    dietary_notes = extracted.get("dietary_notes") or ""

    return PrepContext(
        client_name=client_name,
        destination_label=destination_label,
        cities=cities,
        date_range=date_range,
        hotels=hotels,
        dietary_notes=dietary_notes,
        transport_mode=transport_mode,
    )


def extract_trip_context(input_path: Path, ai_client=None) -> PrepContext:
    """Extract trip context from a DOCX input file or its sibling trip_facts.json.

    If trip_facts.json exists in the same directory as input_path, it is
    loaded and used directly (no AI call needed).  Otherwise the DOCX is
    parsed with parse_itinerary (AI-assisted when ai_client is provided,
    regex-only otherwise).
    """
    facts_path = input_path.parent / "trip_facts.json"
    if facts_path.exists():
        data = json.loads(facts_path.read_text(encoding="utf-8"))
        return _build_prep_context_from_trip_facts(data)

    # For DOCX files, try BVM-specific structural extraction first (handles
    # service vouchers and notes files that the PDF-oriented regex parser
    # cannot read reliably).
    if input_path.suffix.lower() == ".docx":
        extracted = _extract_from_docx(input_path)
        if extracted.get("cities") or extracted.get("client_name"):
            return _build_prep_context_from_extracted(extracted)

    from .parser import parse_itinerary
    itinerary = parse_itinerary(input_path, ai_client=ai_client)
    return _build_prep_context_from_itinerary_data(itinerary)


# ---------------------------------------------------------------------------
# City → country lookup
# ---------------------------------------------------------------------------

def _city_to_country(city: str, db_path: Path) -> str:
    """Return a country/region label for city using _index.json folder coverage."""
    index_path = db_path / "_index.json"
    if not index_path.exists():
        return ""
    index = json.loads(index_path.read_text(encoding="utf-8"))
    folder_coverage: dict[str, list[str]] = index.get("_folder_coverage", {})
    for folder, cities in folder_coverage.items():
        if city in cities:
            return FOLDER_TO_COUNTRY.get(folder, folder)
    return ""


# ---------------------------------------------------------------------------
# Library section renderers
# ---------------------------------------------------------------------------

def _render_restaurants(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Curated Restaurants",
        "*Use these in Day-wise Itinerary restaurant recommendations.*\n",
        "| Name | Cuisine | Area | Hours | Veg-friendly | Must-Try Dishes |",
        "|------|---------|------|-------|:------------:|-----------------|",
    ]
    for r in items:
        name = r.get("name", "")
        cuisine = ", ".join(r.get("cuisine_type") or [])
        area = r.get("area", "")
        hours = r.get("hours", "")
        veg = "Yes" if r.get("vegetarian_friendly") else "No"
        must_try = ", ".join(r.get("must_try_dishes") or [])
        lines.append(f"| {name} | {cuisine} | {area} | {hours} | {veg} | {must_try} |")
    return "\n".join(lines)


def _render_local_dishes(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Must-Try Local Dishes",
        "*Use in the \"Must-Try Local Dishes\" section.*\n",
    ]
    for d in items:
        name = d.get("name", "")
        desc = d.get("description", "")
        where = d.get("where_to_try", "")
        where_str = f" Best at: {where}." if where else ""
        lines.append(f"- **{name}** — {desc}{where_str}")
    return "\n".join(lines)


def _render_souvenirs(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Souvenir Shopping",
        "*Use in the \"Souvenir Shopping Guide\" section.*\n",
    ]
    for s in items:
        item = s.get("item", "")
        where = ", ".join(s.get("where_to_buy") or [])
        where_str = f" — Where to buy: {where}" if where else ""
        lines.append(f"- **{item}**{where_str}")
    return "\n".join(lines)


def _render_transport(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Getting Around",
        "*Use in the \"Getting Around\" section.*\n",
    ]
    for t in items:
        mode = t.get("mode", "") or t.get("type", "")
        desc = t.get("description", "")
        cost = t.get("cost", "")
        cost_str = f" Cost: {cost}" if cost else ""
        lines.append(f"- **{mode}** — {desc}{cost_str}")
    return "\n".join(lines)


def _render_safety(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Safety Tips",
        "*Use in the \"Safety & Emergency Contacts\" section.*\n",
    ]
    for t in items:
        tip = t.get("tip", "") or t.get("description", "")
        if tip:
            lines.append(f"- {tip}")
    return "\n".join(lines)


def _render_connectivity(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Mobile Connectivity",
        "*Use in the \"Mobile Connectivity Guide\" section.*\n",
    ]
    for t in items:
        tip = t.get("tip", "") or t.get("description", "")
        if tip:
            lines.append(f"- {tip}")
    return "\n".join(lines)


def _render_emergency(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Emergency Contacts",
        "*Use in the \"Safety & Emergency Contacts\" section.*\n",
    ]
    for c in items:
        service = c.get("service", "")
        number = c.get("number", "") or c.get("contact", "")
        lines.append(f"- {service}: {number}")
    return "\n".join(lines)


def _render_health(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "### Health & Vaccination",
        "*Use in the \"Health & Vaccination Guidance\" section.*\n",
    ]
    for t in items:
        tip = t.get("tip", "") or t.get("description", "") or t.get("advice", "")
        if tip:
            lines.append(f"- {tip}")
    return "\n".join(lines)


_SECTION_RENDERERS = {
    "restaurants":        _render_restaurants,
    "local_dishes":       _render_local_dishes,
    "souvenirs":          _render_souvenirs,
    "transport_options":  _render_transport,
    "safety_tips":        _render_safety,
    "connectivity_tips":  _render_connectivity,
    "emergency_contacts": _render_emergency,
    "health_tips":        _render_health,
}


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def _render_city_block(city: str, db_path: Path) -> str:
    """Render the full markdown block for one city."""
    country = _city_to_country(city, db_path)
    heading = f"## {city}, {country}" if country else f"## {city}"

    city_file = db_path / f"{city}.json"
    if not city_file.exists():
        return (
            f"{heading}\n\n"
            f"*No BVM library data available for {city} — "
            f"ChatGPT will use general knowledge.*"
        )

    data: dict = json.loads(city_file.read_text(encoding="utf-8"))
    sections: list[str] = [heading]

    for key, _label in _LIBRARY_SECTIONS:
        items = data.get(key) or []
        renderer = _SECTION_RENDERERS.get(key)
        if renderer:
            rendered = renderer(items)
            if rendered:
                sections.append(rendered)

    return "\n\n".join(sections)


def generate_library_context(prep_context: PrepContext, db_path: Path) -> str:
    """Generate the BVM Library Context markdown document."""
    parts = [
        f"# BVM Library Context: {prep_context.destination_label}",
        (
            "> **Instructions for ChatGPT:** The sections below contain BVM's curated\n"
            "> recommendations for this trip. Prioritise these restaurants, dishes, and\n"
            "> facts when generating the guide. Use your own knowledge only to fill gaps\n"
            "> where BVM data is absent or insufficient (fewer than 3 restaurants for a\n"
            "> meal slot, missing a section entirely, etc.)."
        ),
    ]
    for city in prep_context.cities:
        parts.append("---")
        parts.append(_render_city_block(city, db_path))

    return "\n\n".join(parts) + "\n"


def generate_client_profile(prep_context: PrepContext) -> str:
    """Generate the client profile markdown template."""
    ctx = prep_context

    client_line = ctx.client_name if ctx.client_name else "[fill in: client name]"
    title = f"# Client Profile: {client_line} — {ctx.destination_label}"

    destinations_line = ctx.destination_label or "[fill in: destination(s)]"
    dates_line = ctx.date_range if ctx.date_range else "[fill in: travel dates]"

    hotel_lines: list[str] = []
    if ctx.hotels:
        for city, hotel in ctx.hotels.items():
            hotel_lines.append(f"  - {city}: {hotel}")
    else:
        hotel_lines.append("  - [fill in: city: hotel name]")

    diet_line = (
        f"{ctx.dietary_notes} *(extracted from input)*"
        if ctx.dietary_notes
        else "[fill in: e.g. vegetarian / halal / no restrictions]"
    )

    transport_line = (
        f"{ctx.transport_mode} *(noted in input)*"
        if ctx.transport_mode
        else "[fill in: e.g. Public Transport / Taxi / Private car]"
    )

    hotels_block = "\n".join(hotel_lines)

    return f"""{title}

> **Instructions for ChatGPT:** Use this profile to personalise the guide.
> Fields marked [fill in] were not found in the input file — please complete
> these before sending to ChatGPT.

## Trip Details
- **Client Name:** {client_line}
- **Destination(s):** {destinations_line}
- **Travel Dates:** {dates_line}
- **Hotels:**
{hotels_block}

## Dietary Preferences
- **Restrictions / Preferences:** {diet_line}

## Travel Profile
- **Travel Style:** [fill in: budget / mid-range / luxury]
- **Occasion:** [fill in: leisure / honeymoon / anniversary / family / business / other]
- **Group Composition:** [fill in: solo / couple / family with kids / group of friends / other]
- **Transport Mode:** {transport_line}

## Special Requirements
- [fill in: any specific requests, accessibility needs, must-see places, etc.]

---
*Generated automatically from input file. Review and complete [fill in] fields before sending to ChatGPT.*
"""


def _filename_base(prep_context: PrepContext) -> str:
    """Derive a safe filename base from the context, e.g. 'Bhushan_London'."""
    name_part = prep_context.client_name.replace(" ", "_") if prep_context.client_name else "client"
    dest_part = prep_context.cities[0].replace(" ", "_") if prep_context.cities else "trip"
    return f"{name_part}_{dest_part}"


def run_prep(
    input_path: Path,
    db_path: Path,
    output_dir: Optional[Path] = None,
    ai_client=None,
) -> tuple[Path, Path]:
    """Orchestrate prep: extract context, generate docs, write files.

    Returns (library_context_path, client_profile_path).
    """
    ctx = extract_trip_context(input_path, ai_client=ai_client)

    library_md = generate_library_context(ctx, db_path)
    profile_md = generate_client_profile(ctx)

    out_dir = output_dir if output_dir is not None else input_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    base = _filename_base(ctx)
    context_path = out_dir / f"{base}_library_context.md"
    profile_path = out_dir / f"{base}_client_profile.md"

    context_path.write_text(library_md, encoding="utf-8")
    profile_path.write_text(profile_md, encoding="utf-8")

    return context_path, profile_path
