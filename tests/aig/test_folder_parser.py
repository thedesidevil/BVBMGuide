import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from src.aig.folder_parser import FolderParser
from src.common.models import TripFacts


MINIMAL_FACTS_DICT = {
    "client_names": ["Ana"],
    "num_guests": "2 adults",
    "departure_city": "Mumbai",
    "destinations": ["Paris"],
    "trip_start_date": "2025-06-01",
    "trip_end_date": "2025-06-07",
    "hotels": [{"city": "Paris", "hotel_name": "Hotel Paris",
                "check_in": "2025-06-01", "check_out": "2025-06-07"}],
    "transport_modes": [{"from_city": "Mumbai", "to_city": "Paris", "mode": "flight",
                         "date": "2025-06-01", "operator": "Air France", "ticket_number": "AF 225",
                         "departure_time": "02:15", "arrival_time": "08:30",
                         "from_terminal": "Mumbai CSIA T2", "to_terminal": "Paris CDG T2E",
                         "booking_reference": "XYZ789"}],
    "local_transport": "Public Transport",
    "days": [{"day_number": 1, "date": "2025-06-01", "title": "Arrival",
              "overnight_city": "Paris", "overnight_hotel": "Hotel Paris",
              "activities": ["Check-in"],
              "transport_note": "Taxi from airport to hotel"}],
    "dietary_restrictions": [],
    "food_allergies": [],
    "cuisine_preferences": [],
}


@pytest.fixture
def mock_ai():
    client = MagicMock()
    client.complete.return_value = json.dumps(MINIMAL_FACTS_DICT)
    client.complete_with_images.return_value = json.dumps(MINIMAL_FACTS_DICT)
    return client


@pytest.fixture
def parser(mock_ai):
    return FolderParser(ai_client=mock_ai)


class TestFindFiles:
    def test_finds_pdf_and_docx(self, tmp_path, parser):
        (tmp_path / "itinerary.pdf").write_bytes(b"fake")
        (tmp_path / "hotel.docx").write_bytes(b"fake")
        (tmp_path / "notes.txt").write_text("ignore")
        names = {f.name for f in parser._find_files(tmp_path)}
        assert "itinerary.pdf" in names
        assert "hotel.docx" in names
        assert "notes.txt" not in names

    def test_returns_sorted_by_name(self, tmp_path, parser):
        (tmp_path / "z_last.pdf").write_bytes(b"fake")
        (tmp_path / "a_first.pdf").write_bytes(b"fake")
        files = parser._find_files(tmp_path)
        assert files[0].name == "a_first.pdf"
        assert files[1].name == "z_last.pdf"

    def test_empty_folder_returns_empty_list(self, tmp_path, parser):
        assert parser._find_files(tmp_path) == []


class TestExtractText:
    def test_extract_docx_text(self, tmp_path, parser):
        doc = Document()
        doc.add_paragraph("Hello Amsterdam")
        doc.add_paragraph("Day 1: Arrival")
        path = tmp_path / "trip.docx"
        doc.save(str(path))
        text = parser._extract_text(path)
        assert text is not None
        assert "Hello Amsterdam" in text
        assert "Day 1: Arrival" in text

    def test_returns_none_on_corrupt_pdf(self, tmp_path, parser):
        bad = tmp_path / "bad.pdf"
        bad.write_bytes(b"not a real pdf")
        assert parser._extract_text(bad) is None

    def test_returns_none_on_nonexistent_file(self, tmp_path, parser):
        assert parser._extract_text(tmp_path / "ghost.pdf") is None


class TestExtractFactsFromFile:
    def test_returns_parsed_dict(self, parser, mock_ai):
        result = parser._extract_facts_from_file("some travel text", "itinerary.pdf")
        assert result["client_names"] == ["Ana"]
        assert result["destinations"] == ["Paris"]
        mock_ai.complete.assert_called_once()

    def test_handles_json_with_code_fences(self, parser, mock_ai):
        mock_ai.complete.return_value = f"```json\n{json.dumps(MINIMAL_FACTS_DICT)}\n```"
        result = parser._extract_facts_from_file("text", "file.pdf")
        assert result["client_names"] == ["Ana"]

    def test_returns_empty_dict_on_bad_json(self, parser, mock_ai):
        mock_ai.complete.return_value = "not valid json at all {{{"
        result = parser._extract_facts_from_file("text", "file.pdf")
        assert result == {}

    def test_truncates_long_text_to_40000_chars(self, parser, mock_ai):
        long_text = "x" * 100_000
        parser._extract_facts_from_file(long_text, "big.pdf")
        prompt_used = mock_ai.complete.call_args[0][0]
        # Prompt template overhead is ~3 100 chars; text is capped at 40 000.
        # If truncation were removed, prompt would be ~103 100 chars.
        assert len(prompt_used) < 44_000


class TestMergeFacts:
    def test_returns_trip_facts_instance(self, parser, mock_ai):
        partials = [
            ("itinerary.pdf", MINIMAL_FACTS_DICT),
            ("flight.pdf", {"client_names": [], "num_guests": None, "departure_city": "Mumbai",
                            "destinations": [], "trip_start_date": None, "trip_end_date": None,
                            "hotels": [], "transport_modes": [], "days": [],
                            "dietary_restrictions": [], "food_allergies": [], "cuisine_preferences": []}),
        ]
        result = parser._merge_facts(partials)
        assert isinstance(result, TripFacts)
        mock_ai.complete.assert_called_once()

    def test_merge_prompt_includes_all_filenames(self, parser, mock_ai):
        partials = [("a.pdf", MINIMAL_FACTS_DICT), ("b.pdf", MINIMAL_FACTS_DICT)]
        parser._merge_facts(partials)
        prompt = mock_ai.complete.call_args[0][0]
        assert "a.pdf" in prompt
        assert "b.pdf" in prompt

    def test_merge_returns_empty_facts_on_garbage_response(self, parser, mock_ai):
        mock_ai.complete.return_value = "not json at all {{{"
        partials = [("a.pdf", MINIMAL_FACTS_DICT), ("b.pdf", MINIMAL_FACTS_DICT)]
        result = parser._merge_facts(partials)
        assert isinstance(result, TripFacts)
        assert result.client_names == []


class TestParse:
    def test_empty_folder_returns_empty_facts(self, tmp_path, mock_ai):
        p = FolderParser(ai_client=mock_ai)
        facts, warnings = p.parse(tmp_path)
        assert isinstance(facts, TripFacts)
        assert facts.client_names == []
        assert warnings == []
        mock_ai.complete.assert_not_called()

    def test_single_file_extracts_without_merge(self, tmp_path, mock_ai):
        doc = Document()
        doc.add_paragraph("Trip content")
        doc.save(str(tmp_path / "itinerary.docx"))
        p = FolderParser(ai_client=mock_ai)
        facts, warnings = p.parse(tmp_path)
        assert isinstance(facts, TripFacts)
        assert warnings == []
        assert mock_ai.complete.call_count == 1  # extract only, no merge

    def test_multiple_files_calls_merge(self, tmp_path, mock_ai):
        for name in ("a.docx", "b.docx"):
            doc = Document()
            doc.add_paragraph("content")
            doc.save(str(tmp_path / name))
        p = FolderParser(ai_client=mock_ai)
        p.parse(tmp_path)
        assert mock_ai.complete.call_count == 3  # 2 extractions + 1 merge

    def test_unreadable_file_added_to_warnings(self, tmp_path, mock_ai):
        (tmp_path / "bad.pdf").write_bytes(b"not a pdf")
        doc = Document()
        doc.add_paragraph("Good content")
        doc.save(str(tmp_path / "good.docx"))
        p = FolderParser(ai_client=mock_ai)
        facts, warnings = p.parse(tmp_path)
        assert "bad.pdf" in warnings
        assert mock_ai.complete.call_count == 1  # only good.docx extracted

    def test_all_files_unreadable_returns_empty_facts(self, tmp_path, mock_ai):
        (tmp_path / "bad.pdf").write_bytes(b"not a pdf")
        p = FolderParser(ai_client=mock_ai)
        facts, warnings = p.parse(tmp_path)
        assert facts.client_names == []
        assert "bad.pdf" in warnings
        mock_ai.complete.assert_not_called()


class TestRenderPdfPages:
    def test_returns_empty_list_for_corrupt_pdf(self, tmp_path, parser):
        bad = tmp_path / "corrupt.pdf"
        bad.write_bytes(b"not a pdf")
        result = parser._render_pdf_pages(bad)
        assert result == []

    def test_returns_empty_list_for_nonexistent_file(self, tmp_path, parser):
        result = parser._render_pdf_pages(tmp_path / "missing.pdf")
        assert result == []


class TestExtractFactsViaVision:
    def test_returns_none_when_no_pages_rendered(self, parser, mock_ai, monkeypatch):
        monkeypatch.setattr(parser, "_render_pdf_pages", lambda f, **kw: [])
        result = parser._extract_facts_via_vision(Path("ticket.pdf"))
        assert result is None
        mock_ai.complete_with_images.assert_not_called()

    def test_calls_ai_with_images_and_returns_dict(self, parser, mock_ai, monkeypatch):
        monkeypatch.setattr(parser, "_render_pdf_pages", lambda f, **kw: ["base64data"])
        result = parser._extract_facts_via_vision(Path("ticket.pdf"))
        assert result is not None
        assert result["client_names"] == ["Ana"]
        mock_ai.complete_with_images.assert_called_once()
        call_kwargs = mock_ai.complete_with_images.call_args
        assert call_kwargs[1]["images"] == ["base64data"]

    def test_filename_included_in_prompt(self, parser, mock_ai, monkeypatch):
        monkeypatch.setattr(parser, "_render_pdf_pages", lambda f, **kw: ["base64data"])
        parser._extract_facts_via_vision(Path("my_flight_ticket.pdf"))
        prompt = mock_ai.complete_with_images.call_args[0][0]
        assert "my_flight_ticket.pdf" in prompt

    def test_returns_none_on_bad_vision_response(self, parser, mock_ai, monkeypatch):
        monkeypatch.setattr(parser, "_render_pdf_pages", lambda f, **kw: ["base64data"])
        mock_ai.complete_with_images.return_value = "not valid json {{{"
        result = parser._extract_facts_via_vision(Path("ticket.pdf"))
        assert result == {}


class TestParseWithVisionFallback:
    def test_vision_used_for_image_only_pdf(self, tmp_path, mock_ai, monkeypatch):
        (tmp_path / "ticket.pdf").write_bytes(b"fake image pdf")
        p = FolderParser(ai_client=mock_ai)
        monkeypatch.setattr(p, "_render_pdf_pages", lambda f, **kw: ["base64img"])
        facts, warnings = p.parse(tmp_path)
        assert warnings == []
        assert facts.client_names == ["Ana"]
        mock_ai.complete_with_images.assert_called_once()
        mock_ai.complete.assert_not_called()

    def test_warning_when_vision_also_fails(self, tmp_path, mock_ai, monkeypatch):
        (tmp_path / "ticket.pdf").write_bytes(b"fake image pdf")
        p = FolderParser(ai_client=mock_ai)
        monkeypatch.setattr(p, "_render_pdf_pages", lambda f, **kw: [])
        facts, warnings = p.parse(tmp_path)
        assert "ticket.pdf" in warnings
        mock_ai.complete_with_images.assert_not_called()

    def test_docx_does_not_attempt_vision(self, tmp_path, mock_ai, monkeypatch):
        from docx import Document
        doc = Document()
        # Empty paragraph — pdfplumber won't be called but docx text will be blank
        doc.save(str(tmp_path / "empty.docx"))
        p = FolderParser(ai_client=mock_ai)
        vision_called = []
        monkeypatch.setattr(p, "_extract_facts_via_vision", lambda f: vision_called.append(f) or None)
        p.parse(tmp_path)
        assert vision_called == []
