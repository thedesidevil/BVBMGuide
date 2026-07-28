from __future__ import annotations
import io
import re
import urllib.parse
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.hotel_options.models import Plan, EnrichedHotel, HotelRow


# ── Theme system ──────────────────────────────────────────────────────────────

@dataclass
class Theme:
    header_hex: str
    light_hex: str
    marinas_hex: str
    border_hex: str
    details_right_hex: str
    secondary_hex: str
    primary: RGBColor
    secondary: RGBColor
    accent: RGBColor


_THEME_NAVY = Theme(
    header_hex="1F3A5F",
    light_hex="F3F7FB",
    marinas_hex="F3F7FB",
    border_hex="B8C7D9",
    details_right_hex="FAFCFE",
    secondary_hex="5E789A",
    primary=RGBColor(0x1F, 0x3A, 0x5F),
    secondary=RGBColor(0x5E, 0x78, 0x9A),
    accent=RGBColor(0x2E, 0x86, 0xC1),
)

_THEME_GOLD = Theme(
    header_hex="8A6D2F",
    light_hex="FBF6EA",
    marinas_hex="FFFDF7",
    border_hex="E5D6B8",
    details_right_hex="FFFDF7",
    secondary_hex="C77800",
    primary=RGBColor(0x8A, 0x6D, 0x2F),
    secondary=RGBColor(0xC7, 0x78, 0x00),
    accent=RGBColor(0xC7, 0x78, 0x00),
)


def _theme(index: int) -> Theme:
    return _THEME_NAVY if index % 2 == 0 else _THEME_GOLD


# ── Shared constants ──────────────────────────────────────────────────────────

_CHARCOAL     = RGBColor(0x2D, 0x2D, 0x2D)
_GREY         = RGBColor(0x66, 0x66, 0x66)
_GREEN        = RGBColor(0x2E, 0x7D, 0x32)
_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
_NAVY         = RGBColor(0x1F, 0x3A, 0x5F)
_GOLD         = RGBColor(0x8A, 0x6D, 0x2F)
_AMBER        = RGBColor(0xC7, 0x78, 0x00)
_SAVINGS_BG      = "EAF4EA"
_REC_BANNER_BG   = "FBF6EA"
_WHITE_BG        = "FFFFFF"
_RULE_COLOR      = "CCCCCC"
_OUTER_BORDER_HEX = "95B3D7"
_AMBER_BORDER_HEX = "C77800"
_SNAP_CELL_BORDER = "D9D2C3"
_FONT         = "Georgia"
_MARGIN       = Inches(0.75)


# ── Number formatting ─────────────────────────────────────────────────────────

def format_indian_number(amount: float) -> str:
    n = int(round(amount))
    s = str(n)
    if len(s) <= 3:
        return f"₹{s}"
    last3 = s[-3:]
    rest = s[:-3]
    groups: list[str] = []
    while len(rest) > 2:
        groups.insert(0, rest[-2:])
        rest = rest[:-2]
    if rest:
        groups.insert(0, rest)
    return f"₹{','.join(groups)},{last3}"


def _star_category(category: str) -> str:
    """'4-Star Hotel' → '4-star', '3 Star' → '3-star', else ''."""
    m = re.search(r'(\d)', category or "")
    return f"{m.group(1)}-star" if m else ""


# ── Document setup ────────────────────────────────────────────────────────────

def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin   = _MARGIN
        section.right_margin  = _MARGIN


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name  = _FONT
    normal.font.size  = Pt(11)
    normal.font.color.rgb = _CHARCOAL
    normal.paragraph_format.space_after  = Pt(0)
    normal.paragraph_format.left_indent  = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name  = "Georgia"
    h1.font.size  = Pt(20)
    h1.font.bold  = False
    h1.font.color.rgb = _NAVY
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.left_indent  = Pt(0)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _blank(doc: Document, n: int = 1) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        _sp(p, 0, 0)


def _sp(para, before: float = 0, after: float = 0) -> None:
    fmt = para.paragraph_format
    if before:
        fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def _run(para, text: str, *, font: str = _FONT, size: float | None = None,
         bold: bool = False, italic: bool = False,
         color: RGBColor = _CHARCOAL) -> None:
    r = para.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    r.font.name      = font
    if size is not None:
        r.font.size = Pt(size)
    r.font.color.rgb = color


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top: float, left: float,
                      bottom: float, right: float) -> None:
    """Set cell padding in pt (converted to twips internally)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(int(val * 20)))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _pin_table_left(table) -> None:
    """Set explicit tblW + tblInd=0 so the table left-aligns with body text.
    tblInd must precede tblBorders in the schema, so insert at position 1."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "10080")  # 7.0 in × 1440 dxa/in = full text column
    tblW.set(qn("w:type"), "dxa")
    tbl_pr.insert(0, tblW)
    for existing in tbl_pr.findall(qn("w:tblInd")):
        tbl_pr.remove(existing)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tbl_pr.insert(1, tblInd)  # must come before tblBorders per OOXML schema


def _no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def _thin_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), _RULE_COLOR)
                tcBorders.append(b)
            tcPr.append(tcBorders)


def _tbl_outer_borders(table, hex_color: str = "95B3D7", sz: int = 4) -> None:
    """Apply table-level single borders on all sides and inside grid lines."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), hex_color)
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def _cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
    """Set cell-level borders. Each side: 'nil' = no border, (sz, hex) = single, None = skip."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tb = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        b = OxmlElement(f"w:{side}")
        if spec == "nil":
            b.set(qn("w:val"), "nil")
            tb.append(b)
        elif spec is not None:
            sz, color = spec
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            tb.append(b)
    if len(tb):
        tcPr.append(tb)


def _micro_gap(doc: Document) -> None:
    """Near-zero-height separator between adjacent tables; suppresses all spacing."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "1")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)
    # 1pt paragraph-mark font eliminates Word's font-size floor on exact line height
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "2")
    rPr.append(sz)
    rPr.append(szCs)
    pPr.append(rPr)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    _sp(p, 0, 0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _line_spacing_15(para) -> None:
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")


# ── Cover page helpers ────────────────────────────────────────────────────────

def _georgia(para, text: str, size: float, bold: bool = False,
             italic: bool = False, color: RGBColor = _CHARCOAL) -> None:
    _run(para, text, font="Georgia", size=size, bold=bold, italic=italic, color=color)


def _add_trip_snapshot(doc: Document, destination: str, requirements: str,
                       stay_requirements: str = "") -> None:
    import re as _re
    # Split on newlines/commas first, then also on dash-separated list items
    # e.g. "1 room- 2 adults + 1child" → ["1 room", "2 adults + 1child"]
    raw_parts = [r.strip() for r in _re.split(r'[\n,]+', requirements) if r.strip()]
    req_lines = []
    for part in raw_parts:
        req_lines.extend(s.strip() for s in _re.split(r'\s*-\s+', part) if s.strip())
    travellers = next((l for l in req_lines if _re.search(r'\d+\s+adult', l, _re.I)), "")
    rooms      = next((l for l in req_lines if _re.search(r'\d+\s+room', l, _re.I)), "")

    data = [
        ("DESTINATION", destination),
        ("TRAVELLERS",  travellers or "—"),
        ("ROOMS",       rooms or "—"),
        ("PREFERENCES", stay_requirements or "—"),
    ]

    _SNAP_HDR = "1F3A5F"
    _SNAP_ALT = "F7F3EA"
    _SNAP_LBL = RGBColor(0x8A, 0x6D, 0x2F)

    table = doc.add_table(rows=2, cols=4)
    _tbl_outer_borders(table, _OUTER_BORDER_HEX, sz=4)

    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[3])
    _shade_cell(hdr, _SNAP_HDR)
    _set_cell_margins(hdr, 7.2, 7.2, 7.2, 7.2)
    _cell_borders(hdr, (6, _SNAP_HDR), (6, _SNAP_HDR), (6, _SNAP_HDR), (6, _SNAP_HDR))
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(hp, 0, 0)
    _run(hp, "TRIP SNAPSHOT", font="Georgia", size=10, bold=True, color=_WHITE)

    for col_idx, (label, value) in enumerate(data):
        cell = table.rows[1].cells[col_idx]
        bg = _SNAP_ALT if col_idx % 2 == 0 else "FFFFFF"
        _shade_cell(cell, bg)
        _cell_borders(cell, (6, _SNAP_CELL_BORDER), (6, _SNAP_CELL_BORDER),
                      (6, _SNAP_CELL_BORDER), (6, _SNAP_CELL_BORDER))
        _set_cell_margins(cell, 7.2, 7.2, 7.2, 7.2)
        lp = cell.paragraphs[0]
        _sp(lp, 0, 2)
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(lp, label, font="Georgia", size=9, bold=True, color=_SNAP_LBL)
        vp = cell.add_paragraph()
        _sp(vp, 0, 2)
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(vp, value, font="Georgia", size=9, color=_CHARCOAL)


def _add_separator_rule(doc: Document) -> None:
    """Thin horizontal rule — 1×1 full-width table with a navy bottom border."""
    table = doc.add_table(rows=1, cols=1)
    tr = table.rows[0]._tr
    trPr = OxmlElement("w:trPr")
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "1")  # 1 twip — minimal visible height
    trPr.append(trHeight)
    tr.insert(0, trPr)
    cell = table.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    # Bottom border only
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "1F497D")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    # Zero cell margins so border sits flush
    _set_cell_margins(cell, 0, 0, 0, 0)


def _add_advisor_note(doc: Document, destination: str) -> None:
    p = doc.add_paragraph()
    _sp(p, 14, 6)
    _georgia(p, "A NOTE FROM BON VOYAGE BY MARINA", size=12, bold=True, color=_NAVY)

    note = (
        f"Thank you for giving us the opportunity to assist with your "
        f"{destination} journey. "
        f"The options in this document have been carefully reviewed and shortlisted "
        f"based on your preferences, location requirements, flexibility, and overall value. "
        f"We hope this guide helps you choose the stay that is right for you."
    )
    p = doc.add_paragraph()
    _run(p, note, font="Georgia", color=_NAVY)


def _add_letterhead_footer(doc: Document, centered: bool = True) -> None:
    align = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    for text, size, bold, italic in [
        ("Bon Voyage By Marina", 12, True, False),
        ("Bespoke Travel Planning • Premium Stays • Seamless Experiences", 11, False, True),
        ("\U0001f4de +91 86000 15316 | \U0001f4f8 @bonvoyagebymarina | \U0001f310 www.bonvoyagebymarina.com", 11, False, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = align
        _sp(p, 12, 12)
        _run(p, text, font="Georgia", size=size, bold=bold, italic=italic, color=_NAVY)
    p = doc.add_paragraph()
    p.alignment = align
    _sp(p, 12, 12)
    _run(p, "✈️ ", font="Georgia", size=11, color=_NAVY)
    _run(p, "Crafting unforgettable journeys, one trip at a time.", font="Georgia", size=11, italic=True, color=_NAVY)


def _build_cover_page(doc: Document, destination: str, client_name: str,
                      requirements: str, stay_requirements: str = "",
                      destination_photo: bytes | None = None) -> None:
    if destination_photo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 0)
        p.add_run().add_picture(io.BytesIO(destination_photo), width=Inches(7.0))
    else:
        _blank(doc, 4)

    _blank(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 10)
    _georgia(p, f"{destination.upper()} ACCOMMODATION RECOMMENDATIONS", size=26, bold=True, color=_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 18)
    _georgia(p, "Curated by Bon Voyage By Marina", size=13, italic=True, color=_GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 4)
    _georgia(p, "PREPARED EXCLUSIVELY FOR", size=10, color=_GREY)

    if client_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 22)
        _georgia(p, client_name.upper(), size=18, bold=True, color=_NAVY)

    _add_trip_snapshot(doc, destination, requirements, stay_requirements)

    _page_break(doc)
    _add_advisor_note(doc, destination)
    _add_separator_rule(doc)
    _add_letterhead_footer(doc, centered=True)


# ── Thank-you page ────────────────────────────────────────────────────────────

def _build_thank_you_page(doc: Document, destination: str) -> None:
    _blank(doc, 6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(p, 0, 8)
    _run(p, "Thank You", size=20, bold=True)

    for text in [
        (
            f"Thank you for giving Bon Voyage By Marina the opportunity to assist with your {destination} journey. "
            f"We hope the accommodation options in this document help you find the stay that best matches your travel style, preferences, and budget."
        ),
        "Should you wish to explore additional options, alternative locations, upgraded room categories, or other travel arrangements, we would be delighted to assist.",
        f"We look forward to helping create an unforgettable {destination} experience for you.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _sp(p, 0, 6)
        _run(p, text, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(p, 18, 4)
    _run(p, "Warm regards,", size=11)
    _add_letterhead_footer(doc, centered=False)


# ── Hotel card components ─────────────────────────────────────────────────────

def _add_hotel_name_card(doc: Document, hotel_name: str, city: str,
                         category: str, theme: Theme,
                         recommended: bool = False) -> None:
    """Full-width themed name card: hotel name (Georgia 14pt bold) + • City • X-star (10pt)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _no_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    _shade_cell(cell, theme.light_hex)
    _set_cell_margins(cell, 6, 6, 6, 6)
    _cell_borders(cell, (8, theme.header_hex), (8, theme.header_hex),
                  (36, theme.header_hex), (8, theme.header_hex))

    p = cell.paragraphs[0]
    _sp(p, 0, 2)
    r = p.add_run(hotel_name)
    r.font.name      = "Georgia"
    r.font.size      = Pt(14)
    r.font.bold      = True
    r.font.color.rgb = theme.primary

    if recommended:
        rec = p.add_run("  ★ RECOMMENDED")
        rec.font.name      = "Georgia"
        rec.font.size      = Pt(10)
        rec.font.bold      = True
        rec.font.color.rgb = _AMBER

    star = _star_category(category)
    parts = " • ".join(x for x in [city, star] if x)
    if parts:
        r2 = p.add_run(f"  •  {parts}")
        r2.font.size      = Pt(10)
        r2.font.color.rgb = theme.secondary


def _add_hotel_photo(doc: Document, photo_bytes: bytes | None) -> None:
    """Centred hotel photo (5.0") matching the hand-curated reference layout."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 8, 6)
    if photo_bytes:
        p.add_run().add_picture(io.BytesIO(photo_bytes), width=Inches(5.0))
    else:
        _run(p, "[ Photo not available ]", size=9, color=_GREY)


def _add_hotel_details_table(doc: Document, enriched: EnrichedHotel,
                              theme: Theme) -> None:
    """2-col HOTEL DETAILS table: Stay/Rating/Flexibility/Includes left, Room/Address/Phone right."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    _tbl_outer_borders(table, _OUTER_BORDER_HEX, sz=4)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.5)

    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    _shade_cell(hdr, theme.header_hex)
    _set_cell_margins(hdr, 7.2, 7.2, 0, 7.2)
    _cell_borders(hdr, (8, theme.header_hex), (8, theme.header_hex),
                  (8, theme.header_hex), (8, theme.header_hex))
    hp = hdr.paragraphs[0]
    _sp(hp, 0, 0)
    _run(hp, "HOTEL DETAILS", size=10, color=_WHITE)

    left  = table.rows[1].cells[0]
    right = table.rows[1].cells[1]
    _shade_cell(left,  theme.light_hex)
    _shade_cell(right, theme.details_right_hex)
    _set_cell_margins(left,  0, 5.4, 0, 5.4)
    _set_cell_margins(right, 0, 5.4, 0, 5.4)
    _cell_borders(left,  "nil", (8, theme.border_hex), (8, theme.border_hex), (8, theme.border_hex))
    _cell_borders(right, "nil", (8, theme.border_hex), "nil",                 (8, theme.border_hex))

    left_rows = [
        ("Stay",        enriched.dates or "—"),
        ("Rating",      f"{enriched.rating}/5 from {enriched.rating_count:,} reviews"
                        if enriched.rating else "—"),
        ("Flexibility", enriched.cancellation or "—"),
        ("Includes",    enriched.meal_type or "—"),
    ]
    for i, (label, value) in enumerate(left_rows):
        p = left.paragraphs[0] if i == 0 else left.add_paragraph()
        _sp(p, 10 if i == 0 else 0, 0)
        _line_spacing_15(p)
        _run(p, f"{label}: ", size=10, bold=True, color=theme.primary)
        _run(p, value, size=10, color=theme.primary)

    right_rows = [
        ("Room",    enriched.room_type or "—"),
        ("Address", enriched.address or "—"),
        ("Phone",   enriched.phone or "—"),
    ]
    for i, (label, value) in enumerate(right_rows):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        _sp(p, 10 if i == 0 else 0, 0)
        _line_spacing_15(p)
        _run(p, f"{label}: ", size=10, bold=True, color=theme.primary)
        _run(p, value, size=10, color=theme.primary)


def _add_hotel_details_table_unenriched(doc: Document, hotel: HotelRow,
                                         theme: Theme) -> None:
    """Fallback 2-col details table when Google Places enrichment is unavailable."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    _tbl_outer_borders(table, _OUTER_BORDER_HEX, sz=4)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.5)

    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    _shade_cell(hdr, theme.header_hex)
    _set_cell_margins(hdr, 7.2, 7.2, 0, 7.2)
    _cell_borders(hdr, (8, theme.header_hex), (8, theme.header_hex),
                  (8, theme.header_hex), (8, theme.header_hex))
    hp = hdr.paragraphs[0]
    _sp(hp, 0, 0)
    _run(hp, "HOTEL DETAILS", size=10, color=_WHITE)

    left  = table.rows[1].cells[0]
    right = table.rows[1].cells[1]
    _shade_cell(left,  theme.light_hex)
    _shade_cell(right, theme.details_right_hex)
    _set_cell_margins(left,  0, 5.4, 0, 5.4)
    _set_cell_margins(right, 0, 5.4, 0, 5.4)
    _cell_borders(left,  "nil", (8, theme.border_hex), (8, theme.border_hex), (8, theme.border_hex))
    _cell_borders(right, "nil", (8, theme.border_hex), "nil",                 (8, theme.border_hex))

    left_rows = [
        ("Stay",        hotel.dates or "—"),
        ("Flexibility", hotel.cancellation or "—"),
        ("Includes",    hotel.meal_type or "—"),
    ]
    for i, (label, value) in enumerate(left_rows):
        p = left.paragraphs[0] if i == 0 else left.add_paragraph()
        _sp(p, 10 if i == 0 else 0, 0)
        _line_spacing_15(p)
        _run(p, f"{label}: ", size=10, bold=True, color=theme.primary)
        _run(p, value, size=10, color=theme.primary)

    p = right.paragraphs[0]
    _sp(p, 10, 0)
    _line_spacing_15(p)
    _run(p, "Room: ", size=10, bold=True, color=theme.primary)
    _run(p, hotel.room_type or "—", size=10, color=theme.primary)


def _add_marinas_take(doc: Document, description: str, theme: Theme) -> None:
    """Full-width themed description box labelled "Marina's Take:"."""
    if not description:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _no_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    _shade_cell(cell, theme.marinas_hex)
    _set_cell_margins(cell, 6, 6, 6, 6)
    _cell_borders(cell, "nil", "nil", (18, theme.header_hex), "nil")

    p = cell.paragraphs[0]
    _sp(p, 0, 2)
    _run(p, "Our Take: ", size=10, bold=True, italic=True, color=theme.primary)
    _run(p, description, size=10, color=_NAVY)


def _add_why_recommend_hotel_box(doc: Document, why: str, theme: Theme) -> None:
    """Per-hotel 'Why we recommend this hotel' box — grouped layout only."""
    if not why:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _no_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    _shade_cell(cell, theme.light_hex)
    _set_cell_margins(cell, 6, 6, 6, 6)
    _cell_borders(cell, (8, theme.header_hex), (8, theme.header_hex),
                  (8, theme.header_hex), (8, theme.header_hex))

    p = cell.paragraphs[0]
    _sp(p, 0, 0)
    _run(p, "Why we recommend this hotel: ", size=10, bold=True, color=theme.primary)
    _run(p, why, size=10, color=theme.primary)


def _add_hotel_card(doc: Document, row: HotelRow,
                    enriched: EnrichedHotel | None,
                    theme_index: int = 0) -> None:
    """Compose a complete hotel card: name → photo → details → marina's take → why-recommend."""
    theme = _theme(theme_index)
    hotel_name = (enriched.official_name if enriched else None) or row.name
    _add_hotel_name_card(doc, hotel_name, row.city, row.category, theme,
                         recommended=row.recommended)
    if enriched:
        _add_hotel_photo(doc, enriched.photo_bytes)
        _add_hotel_details_table(doc, enriched, theme)
        doc.add_paragraph()
        _add_marinas_take(doc, enriched.description, theme)
    else:
        _add_hotel_details_table_unenriched(doc, row, theme)
    if row.why_recommend:
        _add_why_recommend_hotel_box(doc, row.why_recommend, theme)


# ── Pricing + recommendation boxes ───────────────────────────────────────────

def _add_plan_price_summary(doc: Document, plan: Plan, theme: Theme) -> None:
    """3-col price summary: BEST ONLINE PRICE | OUR PRICE | YOU SAVE (accepts Theme)."""
    pr = plan.pricing
    label = plan.label.upper()
    header_text = f"{label} PRICE SUMMARY"
    if plan.recommended:
        header_text += " • RECOMMENDED"

    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    _tbl_outer_borders(table, _OUTER_BORDER_HEX, sz=4)
    col_w = Inches(7.0 / 3)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_w

    hdr = table.rows[0].cells[0].merge(table.rows[0].cells[2])
    _shade_cell(hdr, theme.header_hex)
    _set_cell_margins(hdr, 7.2, 7.2, 7.2, 7.2)
    _cell_borders(hdr, (6, theme.header_hex), (6, theme.header_hex),
                  (6, theme.header_hex), (6, theme.header_hex))
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(hp, 0, 0)
    _run(hp, header_text, size=10, bold=True, color=_WHITE)

    you_save = pr.customer_discount > 0
    you_save_str = format_indian_number(pr.customer_discount) if you_save else "—"
    you_save_pct = f"{pr.discount_pct:.1f}% off best online prices" if you_save else ""

    cols_data = [
        # (bg, label, label_bold, value, value_size, value_bold, value_color, extra_pct)
        (theme.light_hex, "BEST ONLINE PRICE", True,
         format_indian_number(pr.total_online_price), 18, True, _NAVY, None),
        (_WHITE_BG, "OUR PRICE", True,
         format_indian_number(pr.discounted_price), 18, True, theme.primary, None),
        (_SAVINGS_BG, "YOU SAVE", True,
         you_save_str, 18, True, _GREEN, you_save_pct),
    ]

    for col_idx, (bg, lbl, lbl_bold, val, val_sz, val_bold, val_color, extra) in enumerate(cols_data):
        cell = table.rows[1].cells[col_idx]
        _shade_cell(cell, bg)
        _set_cell_margins(cell, 7.2, 7.2, 7.2, 7.2)
        _cell_borders(cell, (6, theme.border_hex), (6, theme.border_hex),
                      (6, theme.border_hex), (6, theme.border_hex))

        lp = cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_color = _GREEN if col_idx == 2 else (_NAVY if col_idx == 0 else theme.primary)
        _sp(lp, 0, 3)
        _run(lp, lbl, size=9, bold=lbl_bold, color=lbl_color)

        vp = cell.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(vp, 0, 0)
        _run(vp, val, size=val_sz, bold=val_bold, color=val_color)

        if extra:
            ep = cell.add_paragraph()
            ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sp(ep, 0, 0)
            _run(ep, extra, size=9, bold=True, color=_GREEN)


def _add_price_summary(doc: Document, plan: Plan, theme_index: int = 0) -> None:
    """3-col price summary: BEST ONLINE PRICE | OUR PRICE | YOU SAVE (accepts theme_index)."""
    _add_plan_price_summary(doc, plan, _theme(theme_index))


def _add_hotel_price_strip(doc: Document, hotel: HotelRow, theme: Theme) -> None:
    """Headerless 3-col price strip for a single hotel (no-plans layout)."""
    you_save = hotel.customer_discount > 0
    you_save_str = format_indian_number(hotel.customer_discount) if you_save else "—"
    you_save_pct = f"{hotel.discount_pct:.1f}% off best online prices" if you_save else ""
    our_price = hotel.discounted_price if hotel.discounted_price > 0 else hotel.online_price

    cols_data = [
        (theme.light_hex, "BEST ONLINE PRICE", format_indian_number(hotel.online_price), _NAVY),
        (_WHITE_BG,        "OUR PRICE",         format_indian_number(our_price),          theme.primary),
        (_SAVINGS_BG,      "YOU SAVE",           you_save_str,                             _GREEN),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    _tbl_outer_borders(table, _OUTER_BORDER_HEX, sz=4)
    col_w = Inches(7.0 / 3)
    for cell in table.rows[0].cells:
        cell.width = col_w

    for col_idx, (bg, lbl, val, val_color) in enumerate(cols_data):
        cell = table.rows[0].cells[col_idx]
        _shade_cell(cell, bg)
        _set_cell_margins(cell, 7.2, 7.2, 7.2, 7.2)
        _cell_borders(cell, (6, theme.border_hex), (6, theme.border_hex),
                      (6, theme.border_hex), (6, theme.border_hex))

        lbl_color = _GREEN if col_idx == 2 else (_NAVY if col_idx == 0 else theme.primary)
        lp = cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(lp, 0, 3)
        _run(lp, lbl, size=9, bold=True, color=lbl_color)

        vp = cell.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(vp, 0, 0)
        _run(vp, val, size=18, bold=True, color=val_color)

        if col_idx == 2 and you_save_pct:
            ep = cell.add_paragraph()
            ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sp(ep, 0, 0)
            _run(ep, you_save_pct, size=9, bold=True, color=_GREEN)


def _add_recommended_choice_banner(doc: Document, plan_label: str,
                                   why_text: str) -> None:
    """Cream banner shown before the exec summary table when a plan is recommended."""
    if not why_text:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _no_borders(table)
    _pin_table_left(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    _shade_cell(cell, _THEME_NAVY.light_hex)
    _set_cell_margins(cell, 6, 6, 6, 6)
    _cell_borders(cell, (8, _THEME_NAVY.header_hex), (8, _THEME_NAVY.header_hex),
                  (8, _THEME_NAVY.header_hex), (8, _THEME_NAVY.header_hex))

    p = cell.paragraphs[0]
    _sp(p, 0, 0)
    _run(p, f"Recommended choice: {plan_label}. ", bold=True, color=_NAVY)
    _run(p, why_text, color=_NAVY)


def _add_why_recommend_box(doc: Document, plan: Plan) -> None:
    """Transition box placed before a plan's heading."""
    if not plan.why_recommend:
        return
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _no_borders(table)
    _pin_table_left(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    _shade_cell(cell, _REC_BANNER_BG)
    _set_cell_margins(cell, 6, 6, 6, 6)
    _cell_borders(cell, (8, _AMBER_BORDER_HEX), (8, _AMBER_BORDER_HEX),
                  (8, _AMBER_BORDER_HEX), (8, _AMBER_BORDER_HEX))

    p = cell.paragraphs[0]
    _sp(p, 0, 0)
    _run(p, f"Why we recommend {plan.label}: ", size=10, bold=True,
         color=_NAVY)
    _run(p, plan.why_recommend, size=10, color=_NAVY)


# ── Executive summary ─────────────────────────────────────────────────────────

def _build_exec_summary_by_plan(doc: Document, plans: list[Plan], enriched_map: dict) -> None:
    col_widths = [Inches(0.74), Inches(3.32), Inches(1.07), Inches(0.94), Inches(0.93)]
    col_labels = ["Plan", "Hotels", "Best Online Price", "Our Price", "You Save"]
    _HDR = _THEME_NAVY.header_hex

    table = doc.add_table(rows=1 + len(plans), cols=5)
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    hdr_row = table.rows[0]
    tr = hdr_row._tr
    trPr = OxmlElement("w:trPr")
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(35 * 20)))  # 35pt in twips
    trPr.append(trHeight)
    tr.insert(0, trPr)

    for i, label in enumerate(col_labels):
        cell = hdr_row.cells[i]
        _shade_cell(cell, _HDR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 0)
        _run(p, label, size=10, bold=True, color=_WHITE)

    for row_idx, plan in enumerate(plans):
        row = table.rows[row_idx + 1]
        bg = "FFFFFF" if row_idx % 2 == 0 else "F7F7F7"
        t = _theme(row_idx)
        pr = plan.pricing

        you_save_str = (format_indian_number(pr.customer_discount)
                        if pr.customer_discount > 0 else "—")
        you_save_pct = (f"({pr.discount_pct:.1f}% off)"
                        if pr.customer_discount > 0 else None)

        # Plan column
        pc = row.cells[0]
        _shade_cell(pc, bg)
        pc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = pc.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(pp, 0, 0)
        _run(pp, plan.label, size=10, bold=True, color=_NAVY)
        if plan.recommended:
            rp = pc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sp(rp, 2, 0)
            _run(rp, "★ RECOMMENDED", size=10, bold=True, color=_AMBER)

        # Hotels column
        hc = row.cells[1]
        _shade_cell(hc, bg)
        hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for h_idx, hotel in enumerate(plan.hotels):
            hp = hc.paragraphs[0] if h_idx == 0 else hc.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = hp.paragraph_format
            fmt.space_before      = Pt(10)
            fmt.space_after       = Pt(10 if h_idx == len(plan.hotels) - 1 else 0)
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            star = _star_category(hotel.category)
            suffix = f" ({star[0]}*)" if star else ""
            _enriched = enriched_map.get(hotel.name)
            h_name = (_enriched.official_name if _enriched and _enriched.official_name else None) or hotel.name
            _run(hp, f"❖ {h_name}{suffix}", size=10, color=_NAVY)

        # Price columns
        for col_idx, (val, bold, color) in enumerate([
            (format_indian_number(pr.total_online_price), True,  _NAVY),
            (format_indian_number(pr.discounted_price),   True,  _NAVY),
            (you_save_str,                                True,  _GREEN),
        ]):
            cell = row.cells[2 + col_idx]
            _shade_cell(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _sp(p, 0, 0)
            _run(p, val, size=10, bold=bold, color=color)
            if col_idx == 2 and you_save_pct:
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _sp(p2, 0, 2)
                _run(p2, you_save_pct, size=10, bold=True, color=_GREEN)

    _thin_borders(table)


def _build_exec_summary_by_hotel(doc: Document, plans: list[Plan], enriched_map: dict) -> None:
    col_widths = [Inches(1.55), Inches(2.55), Inches(0.9), Inches(0.9), Inches(1.1)]
    col_labels = ["City / Dates", "Hotel", "Online Price", "Our Price", "You Save"]
    _HDR = _THEME_NAVY.header_hex

    hotel_rows = [(plan.label, hotel) for plan in plans for hotel in plan.hotels]

    table = doc.add_table(rows=1 + len(hotel_rows), cols=5)
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    hdr_row = table.rows[0]
    tr = hdr_row._tr
    trPr = OxmlElement("w:trPr")
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(35 * 20)))  # 35pt in twips
    trPr.append(trHeight)
    tr.insert(0, trPr)

    for i, label in enumerate(col_labels):
        cell = hdr_row.cells[i]
        _shade_cell(cell, _HDR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 0)
        _run(p, label, size=10, bold=True, color=_WHITE)

    row_idx = 0
    for plan in plans:
        section_label = plan.label
        for hotel_idx, hotel in enumerate(plan.hotels):
            bg = "FFFFFF" if row_idx % 2 == 0 else "F7F7F7"

            row = table.rows[row_idx + 1]
            row_idx += 1

            our_price = hotel.discounted_price if hotel.discounted_price > 0 else hotel.online_price
            you_save_str = (format_indian_number(hotel.customer_discount)
                            if hotel.customer_discount > 0 else "—")
            you_save_pct = (f"({hotel.discount_pct:.1f}% off)"
                            if hotel.customer_discount > 0 else None)

            star = _star_category(hotel.category)
            suffix = f" ({star[0]}*)" if star else ""
            _enriched = enriched_map.get(hotel.name)
            h_name = (_enriched.official_name if _enriched and _enriched.official_name else None) or hotel.name

            for col_idx, (text, align, color, bold) in enumerate([
                (section_label,                            WD_ALIGN_PARAGRAPH.LEFT,   _NAVY,  False),
                (h_name + suffix,                          WD_ALIGN_PARAGRAPH.LEFT,   _NAVY,  False),
                (format_indian_number(hotel.online_price), WD_ALIGN_PARAGRAPH.CENTER, _NAVY,  False),
                (format_indian_number(our_price),          WD_ALIGN_PARAGRAPH.CENTER, _NAVY,  True),
                (you_save_str,                             WD_ALIGN_PARAGRAPH.CENTER, _GREEN, True),
            ]):
                cell = row.cells[col_idx]
                _shade_cell(cell, bg)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = align
                _sp(p, 2, 2)
                _run(p, text, size=9, bold=bold, color=color)
                if col_idx == 1 and hotel.recommended:
                    rp = cell.add_paragraph()
                    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _sp(rp, 0, 2)
                    _run(rp, "★ RECOMMENDED", size=8, bold=True, color=_AMBER)
                if col_idx == 4 and you_save_pct:
                    p2 = cell.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _sp(p2, 0, 2)
                    _run(p2, you_save_pct, size=8, color=_GREEN)

    _thin_borders(table)


def _build_executive_summary(doc: Document, plans: list[Plan],
                              grouped_by_sections: bool = False,
                              enriched_map: dict | None = None) -> None:
    if not plans:
        return

    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.left_indent = Pt(0)
    r = p.add_run("Executive Summary")
    r.font.name = "Georgia"

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(0)
    _run(p, (
        "A curated overview of all accommodation options, organised by destination and date."
        if grouped_by_sections else
        "A client-ready comparison of both accommodation plans, "
        "highlighting value, comfort, and the recommended option."
    ), color=_NAVY)
    _sp(p, 0, 8)

    em = enriched_map or {}
    if grouped_by_sections:
        _build_exec_summary_by_hotel(doc, plans, em)
    else:
        _build_exec_summary_by_plan(doc, plans, em)
        rec_plan = next((pl for pl in plans if pl.recommended), None)
        if rec_plan and rec_plan.why_recommend:
            p = doc.add_paragraph()
            _sp(p, 4, 0)
            _add_recommended_choice_banner(doc, rec_plan.label, rec_plan.why_recommend)


# ── Plans layout (Task 6) ─────────────────────────────────────────────────────

def _build_plans_layout(doc: Document, plans: list[Plan],
                        enriched_map: dict[str, EnrichedHotel],
                        destination: str) -> None:
    for plan_idx, plan in enumerate(plans):
        t = _theme(plan_idx)

        # Plan heading
        p = doc.add_paragraph(style="Heading 1")
        r = p.add_run(plan.label.upper())
        r.font.name = "Georgia"
        if plan.recommended:
            r2 = p.add_run("  ★ RECOMMENDED")
            r2.font.name      = "Georgia"
            r2.font.size      = Pt(12.5)
            r2.font.color.rgb = t.accent

        # Horizontal rule under plan heading
        p = doc.add_paragraph()
        _sp(p, 0, 4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "18")
        top_border.set(qn("w:space"), "0")
        top_border.set(qn("w:color"), t.secondary_hex)
        pBdr.append(top_border)
        pPr.append(pBdr)

        # Hotels: no break before first (stays on same page as heading); break between hotels
        for hotel_idx, hotel in enumerate(plan.hotels):
            if hotel_idx > 0:
                _page_break(doc)
            enriched = enriched_map.get(hotel.name)
            _add_hotel_card(doc, hotel, enriched, theme_index=plan_idx)

        # Price summary on its own page
        _page_break(doc)
        _add_price_summary(doc, plan, theme_index=plan_idx)

        if plan_idx < len(plans) - 1:
            _page_break(doc)


def _build_grouped_sections(doc: Document, plans: list[Plan],
                             enriched_map: dict[str, EnrichedHotel],
                             destination: str) -> None:
    for section_idx, plan in enumerate(plans):
        if section_idx > 0:
            _page_break(doc)

        t = _theme(section_idx)

        # Section heading — styled like plan headings with horizontal rule
        p = doc.add_paragraph(style="Heading 1")
        p.paragraph_format.left_indent = Pt(0)
        r = p.add_run(plan.label.upper())
        r.font.name = "Georgia"
        r.font.color.rgb = t.primary

        p = doc.add_paragraph()
        _sp(p, 0, 4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "18")
        top_border.set(qn("w:space"), "0")
        top_border.set(qn("w:color"), t.secondary_hex)
        pBdr.append(top_border)
        pPr.append(pBdr)
        for hotel_idx, hotel in enumerate(plan.hotels):
            if hotel_idx > 0:
                _page_break(doc)
            enriched = enriched_map.get(hotel.name)
            h_name = (enriched.official_name if enriched else None) or hotel.name
            h_city = hotel.city or re.sub(r'\s*\(.*\)\s*$', '', plan.label).strip() or destination
            h_cat  = (enriched.category if enriched else None) or hotel.category

            _add_hotel_name_card(doc, h_name, h_city, h_cat, t,
                                 recommended=hotel.recommended)
            if enriched:
                _add_hotel_photo(doc, enriched.photo_bytes)
                _add_hotel_details_table(doc, enriched, t)
                p = doc.add_paragraph()
                _sp(p, 4, 0)
                _add_marinas_take(doc, enriched.description, t)
            else:
                _add_hotel_details_table_unenriched(doc, hotel, t)
            _micro_gap(doc)
            _add_hotel_price_strip(doc, hotel, t)
            if hotel.why_recommend:
                p = doc.add_paragraph()
                _sp(p, 6, 0)
                _add_why_recommend_hotel_box(doc, hotel.why_recommend, t)


def _fix_compat_settings(doc: Document) -> None:
    """Set compatibilityMode=15 (Word 2013+) and remove useFELayout.

    python-docx's default template ships with compatibilityMode=14 (Word 2010)
    and useFELayout. Word 2010 compat mode shifts paragraph indentation relative
    to tables, causing heading/subtitle text to appear right of the table border.
    Mode 15 matches the hand-curated reference document.
    """
    settings_elem = doc.settings.element
    compat_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ms_uri = "http://schemas.microsoft.com/office/word"

    # Remove useFELayout
    fe = settings_elem.find(qn("w:useFELayout"))
    if fe is not None:
        settings_elem.remove(fe)

    # Set compatibilityMode to 15
    compat = settings_elem.find(qn("w:compat"))
    if compat is not None:
        for cs in compat.findall(qn("w:compatSetting")):
            if cs.get(qn("w:name")) == "compatibilityMode":
                cs.set(qn("w:val"), "15")
                break


def build_document(
    plans: list[Plan],
    enriched_map: dict[str, EnrichedHotel],
    client_name: str,
    destination: str,
    requirements: str = "",
    destination_photo: bytes | None = None,
    stay_requirements: str = "",
    grouped_by_sections: bool = False,
) -> bytes:
    doc = Document()
    _set_margins(doc)
    _configure_styles(doc)
    _fix_compat_settings(doc)

    _build_cover_page(doc, destination, client_name, requirements,
                      stay_requirements=stay_requirements,
                      destination_photo=destination_photo)
    _page_break(doc)

    _build_executive_summary(doc, plans, grouped_by_sections=grouped_by_sections,
                             enriched_map=enriched_map)
    _page_break(doc)

    if grouped_by_sections:
        _build_grouped_sections(doc, plans, enriched_map, destination)
    else:
        _build_plans_layout(doc, plans, enriched_map, destination)

    _page_break(doc)
    _build_thank_you_page(doc, destination)

    # Pin every table to the left text margin so they align with body paragraphs
    for tbl_elem in doc.element.body.findall(qn("w:tbl")):
        tblPr = tbl_elem.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        for existing in tblPr.findall(qn("w:tblW")):
            tblPr.remove(existing)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "10080")
        tblW.set(qn("w:type"), "dxa")
        tblPr.insert(0, tblW)
        for existing in tblPr.findall(qn("w:tblInd")):
            tblPr.remove(existing)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "0")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.insert(1, tblInd)  # must come before tblBorders per OOXML schema

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
