"""AIG verification service — two-layer pipeline (rule engine + AI pass)."""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field, asdict

from docx import Document

from src.common.ai_provider import get_ai_client, AIClient
from src.common.doc_extractor import extract_from_text, VERIFY_EXTRACTION_PROMPT


# ---------------------------------------------------------------------------
# Data types
# ---------------------------------------------------------------------------

@dataclass
class Finding:
    check_id: str
    layer: str          # "rule" | "ai"
    severity: str       # "RED" | "YELLOW"
    section: str
    description: str
    evidence: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class VerifyResult:
    findings: list[Finding]
    narratives: dict
    meta: dict

    def to_dict(self) -> dict:
        return {
            "findings": [f.to_dict() for f in self.findings],
            "narratives": self.narratives,
            "meta": self.meta,
        }


# ---------------------------------------------------------------------------
# DOCX text extraction
# ---------------------------------------------------------------------------

def extract_paragraphs(docx_bytes: bytes) -> list[dict]:
    """Return list of {style, text} dicts from DOCX, skipping empty paragraphs."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({"style": para.style.name, "text": text})
    return paragraphs


def paragraphs_to_text(paragraphs: list[dict]) -> str:
    return "\n".join(p["text"] for p in paragraphs)


# ---------------------------------------------------------------------------
# Rule engine (R1–R10)
# ---------------------------------------------------------------------------

_AI_ARTIFACT_PATTERNS = [
    r"^here(?:'s| is) day \d",                              # start-of-line only
    r"sure[!,]?\s+here(?:'s| is)",
    r"i have generated",
    r"i['']ve generated",
    r"\bas an ai,\b",                                        # "as an AI, I ..." not "act as an AI"
    r"\bas an ai assistant\b",
    r"i['']ll (?:provide|create|generate|write|now write)",
    r"here(?:'s| is) the (?:complete|full|detailed|updated|revised)",
    r"^note:\s+this (?:guide|section|day|itinerary)",
    r"```",
    r"^in conclusion,",
]

_PLACEHOLDER_PATTERNS = [
    (r"\[hotel name\]", "hotel name placeholder"),
    (r"\[restaurant name\]", "restaurant name placeholder"),
    (r"\[insert [^\]]+\]", "INSERT placeholder"),
    (r"\[tbd\]", "TBD placeholder"),
    (r"\btbd\b", "TBD"),
    (r"\bxxxxx+\b", "XXXX placeholder"),
    (r"\[(?:city|destination|attraction|place) name\]", "destination placeholder"),
]

_MANDATORY_SECTIONS = [
    "Client Information",
    "Important Places Around Your Stay",
    "Souvenir Shopping Guide",
    "Must-Try Local Dishes",
    "Getting Around",
    "Cultural Etiquette",
    "Tailored Packing List",
    "Mobile Connectivity",
    "Safety & Emergency Contacts",
    "Health & Vaccination",
    "Thank You",
]

_WEEKDAYS = [
    "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday",
    "Mon,", "Tue,", "Wed,", "Thu,", "Fri,", "Sat,", "Sun,",
]


def _r1_ai_artifacts(text: str, findings: list[Finding]) -> None:
    for pattern in _AI_ARTIFACT_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if m:
            start = max(0, m.start() - 20)
            end = min(len(text), m.end() + 60)
            findings.append(Finding(
                check_id="R1",
                layer="rule",
                severity="RED",
                section="Entire document",
                description=f"AI meta-commentary or artifact text detected in document",
                evidence=text[start:end].strip(),
            ))
            return  # one R1 finding is enough


def _r2_placeholders(text: str, findings: list[Finding]) -> None:
    for pattern, label in _PLACEHOLDER_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            findings.append(Finding(
                check_id="R2",
                layer="rule",
                severity="RED",
                section="Entire document",
                description=f"Unfilled placeholder text found: {label}",
                evidence=m.group(),
            ))
            return


def _r3_mandatory_sections(text: str, findings: list[Finding]) -> None:
    for section in _MANDATORY_SECTIONS:
        if not re.search(re.escape(section), text, re.IGNORECASE):
            findings.append(Finding(
                check_id="R3",
                layer="rule",
                severity="RED",
                section="Document structure",
                description=f"Mandatory section missing: '{section}'",
                evidence="",
            ))


def _r4_r10_day_numbers(text: str, findings: list[Finding]) -> None:
    # Extract all day numbers from headings like "Day 1:", "Day 2 –", "Day 3 |"
    day_nums = [int(m.group(1)) for m in re.finditer(r"^Day (\d+)\s*[:|\-–]", text, re.MULTILINE)]
    if not day_nums:
        return
    unique_nums = sorted(set(day_nums))
    # Check for duplicates
    if len(day_nums) != len(unique_nums):
        from collections import Counter
        dupes = [n for n, c in Counter(day_nums).items() if c > 1]
        findings.append(Finding(
            check_id="R10",
            layer="rule",
            severity="RED",
            section="Day headings",
            description=f"Duplicate day number(s) found: Day {', '.join(str(d) for d in sorted(dupes))}",
            evidence=f"Day numbers found: {day_nums[:20]}",
        ))
    # Check for gaps (ignoring Day 0 as a pre-trip day)
    start = min(unique_nums)
    expected = list(range(start, max(unique_nums) + 1))
    missing = [n for n in expected if n not in unique_nums]
    if missing:
        findings.append(Finding(
            check_id="R4",
            layer="rule",
            severity="RED",
            section="Day-wise itinerary",
            description=f"Day number gap(s) detected: Day {', '.join(str(n) for n in missing)} missing",
            evidence=f"Day numbers present: {unique_nums}",
        ))


_WEEKDAY_RE = re.compile(
    r"\b(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday|"
    r"Mon,|Tue,|Wed,|Thu,|Fri,|Sat,|Sun,)\b",
    re.IGNORECASE,
)


def _r5_day_heading_format(text: str, findings: list[Finding]) -> None:
    missing_weekday = []
    for m in re.finditer(r"^(Day \d+\s*[:|\-–].+)$", text, re.MULTILINE):
        heading = m.group(1)
        if not _WEEKDAY_RE.search(heading):
            missing_weekday.append(heading[:80])
    if missing_weekday:
        findings.append(Finding(
            check_id="R5",
            layer="rule",
            severity="YELLOW",
            section="Day headings",
            description=f"{len(missing_weekday)} day heading(s) are missing the day of week (e.g. 'Monday,')",
            evidence=missing_weekday[0],
        ))


def _count_maps_hyperlinks(docx_bytes: bytes) -> int:
    """Count Google Maps URLs stored as hyperlink relationships in the DOCX XML.

    python-docx paragraph.text never includes hyperlink URLs — they live in the
    OPC relationship store (part.rels), so plain-text scanning always returns 0.
    """
    doc = Document(io.BytesIO(docx_bytes))
    _MAPS_RE = re.compile(r"maps\.google\.com|maps\.app\.goo\.gl|goo\.gl/maps", re.IGNORECASE)
    count = sum(1 for r in doc.part.rels.values() if _MAPS_RE.search(r.target_ref))
    # Also check header/footer parts
    for part in doc.part.package.iter_parts():
        try:
            count += sum(1 for r in part.rels.values() if _MAPS_RE.search(r.target_ref))
        except Exception:
            pass
    return count


def _r6_maps_links(text: str, findings: list[Finding], maps_count: int = 0) -> None:
    count = maps_count or len(re.findall(r"maps\.google\.com|maps\.app\.goo\.gl|goo\.gl/maps", text, re.IGNORECASE))
    if count == 0:
        findings.append(Finding(
            check_id="R6",
            layer="rule",
            severity="YELLOW",
            section="Entire document",
            description="No Google Maps links found. Hotels, restaurants, and attractions should each have a Maps link.",
            evidence="",
        ))
    elif count < 5:
        findings.append(Finding(
            check_id="R6",
            layer="rule",
            severity="YELLOW",
            section="Entire document",
            description=f"Very few Maps links found ({count} total). Expected links on hotels, restaurants, and attractions.",
            evidence=f"Only {count} Maps link(s) in entire document",
        ))


def _r7_restaurant_count(text: str, findings: list[Finding]) -> None:
    # Split by day headings
    day_blocks = re.split(r"(?=^Day \d+\s*[:|\-–])", text, flags=re.MULTILINE)
    for block in day_blocks:
        day_match = re.match(r"^(Day \d+)", block)
        if not day_match:
            continue
        day_label = day_match.group(1)
        # Skip cruise/transit days
        if re.search(r"at sea|cruise day|in transit", block, re.IGNORECASE):
            continue
        # Count 🍴 restaurant entries
        restaurant_count = len(re.findall(r"🍴", block))
        # Only flag if day has some content but too few restaurants
        if 0 < restaurant_count < 3:
            findings.append(Finding(
                check_id="R7",
                layer="rule",
                severity="YELLOW",
                section=day_label,
                description=f"{day_label} has only {restaurant_count} restaurant recommendation(s); minimum 3 expected",
                evidence=f"Found {restaurant_count} 🍴 entries",
            ))


_TIME_RANGE_RE = re.compile(
    r"\b(\d{1,2}:\d{2})\s*(AM|PM)?\s*[–\-]\s*(\d{1,2}:\d{2})\s*(AM|PM)\b",
    re.IGNORECASE,
)
# Captures only the closing side of a time range (groups 1=hour, 2=min or None, 3=period)
# Handles both "HH:MM AM/PM" and "H AM/PM" (no minutes) on either side
_CLOSING_TIME_RE = re.compile(
    r"\b\d{1,2}(?::\d{2})?\s*(?:AM|PM)\s*[–\-]\s*(\d{1,2})(?::(\d{2}))?\s*(AM|PM)\b",
    re.IGNORECASE,
)
# Captures only the opening side of a time range (groups 1=hour, 2=min or None, 3=period)
_OPENING_TIME_RE = re.compile(
    r"\b(\d{1,2})(?::(\d{2}))?\s*(AM|PM)\s*[–\-]\s*\d{1,2}(?::\d{2})?\s*(?:AM|PM)\b",
    re.IGNORECASE,
)

_MIDNIGHT_START_RE = re.compile(r"\b12:\d{2}\s*AM\s*[–\-]", re.IGNORECASE)
_SAME_TIME_RE = re.compile(
    r"\b(\d{1,2}:\d{2}\s*(?:AM|PM))\s*[–\-]\s*(\d{1,2}:\d{2}\s*(?:AM|PM))\b",
    re.IGNORECASE,
)
_ALL_EMOJI_RE = re.compile(r"[\U00010000-\U0010FFFF☀-⟿️‍]", re.UNICODE)


def _venue_name_before(text: str, pos: int) -> str:
    """Return the nearest venue-name line before pos (used by R8/R12).

    Strips emojis, skips label lines (contain a colon), prose sentences, and
    long description lines. Does not require a capital letter — venue names can
    start with a number (e.g. '3 Coins') or lowercase letter.
    """
    window = text[max(0, pos - 400):pos]
    for line in reversed(window.splitlines()):
        clean = _ALL_EMOJI_RE.sub("", line).strip()
        if not clean:
            continue
        if not any(c.isalnum() for c in clean):
            continue
        if ":" in clean[:50]:
            continue
        if re.match(r"^(A |An |The |Approx)", clean):
            continue
        if clean.endswith(".") or clean.endswith("!"):
            continue
        if len(clean) > 80:
            continue
        return clean
    return ""


def _r8_time_format(text: str, findings: list[Finding]) -> None:
    # Check 1: first time missing AM/PM when second has it ("12:00 – 11:00 PM")
    for m in _TIME_RANGE_RE.finditer(text):
        if not m.group(2):  # first time has no AM/PM
            full_match = m.group()
            venue = _venue_name_before(text, m.start())
            desc = f"Incomplete time format: '{full_match}' — opening time is missing AM/PM"
            if venue:
                desc = f"{venue}: {desc}"
            findings.append(Finding(
                check_id="R8",
                layer="rule",
                severity="RED",
                section="Opening hours",
                description=desc,
                evidence=full_match,
            ))
            break  # one finding of this type is enough

    # Check 2: midnight start like "12:30 AM – 10:30 PM" — almost always a typo for 12:30 PM
    for m in _MIDNIGHT_START_RE.finditer(text):
        venue = _venue_name_before(text, m.start())
        snippet = text[max(0, m.start() - 10):m.end() + 30].strip()
        desc = f"Opening time starts at midnight — almost certainly a typo for PM"
        if venue:
            desc = f"{venue}: opening time '{m.group().strip()}' starts at midnight — almost certainly a typo for PM"
        findings.append(Finding(
            check_id="R8",
            layer="rule",
            severity="RED",
            section="Opening hours",
            description=desc,
            evidence=snippet,
        ))
        break

    # Check 3: identical start and end time ("11:30 PM – 11:30 PM") — zero-duration or typo
    for m in _SAME_TIME_RE.finditer(text):
        t1 = re.sub(r"\s+", " ", m.group(1)).strip().upper()
        t2 = re.sub(r"\s+", " ", m.group(2)).strip().upper()
        if t1 == t2:
            venue = _venue_name_before(text, m.start())
            desc = f"Opening hours '{m.group().strip()}' have identical start and end time — likely a typo (e.g. 11:30 AM – 11:30 PM intended)"
            if venue:
                desc = f"{venue}: opening hours '{m.group().strip()}' have identical start and end time — likely a typo (e.g. 11:30 AM – 11:30 PM intended)"
            findings.append(Finding(
                check_id="R8",
                layer="rule",
                severity="RED",
                section="Opening hours",
                description=desc,
                evidence=m.group(),
            ))
            break


_ENC_RE = re.compile(r'[�-￾﻿]')


def _to_minutes(hour: int, minute: int | None, period: str) -> int:
    """Convert 12-hour time to minutes since midnight. minute may be None if not present."""
    m = minute or 0
    p = period.upper()
    if p == "PM" and hour != 12:
        return (hour + 12) * 60 + m
    if p == "AM" and hour == 12:
        return m
    return hour * 60 + m


def _r11_meal_timing(meal_venues: list[dict], findings: list[Finding]) -> None:
    """Check meal venue opening hours against meal-time thresholds.

    Uses AI-extracted structured data (name, meal_section, opening_hours) so
    venue names and hours are reliable regardless of AIG formatting conventions.

    Thresholds:
      Dinner  — closing time must be ≥ 9 PM (21:00)
      Lunch   — closing time must be ≥ 2 PM (14:00)
      Breakfast — opening time must be ≤ 8 AM (08:00)

    For split-hours venues (e.g. "11 AM–3 PM | 5 PM–10 PM") the best session
    is used: latest close for dinner/lunch, earliest open for breakfast.
    """
    configs = {
        "Dinner":    {"check": "closing",  "threshold": 21 * 60, "label": "dinner recommendation — dinner venues should be open until at least 9 PM"},
        "Lunch":     {"check": "closing",  "threshold": 14 * 60, "label": "lunch recommendation — lunch venues should be open until at least 2 PM"},
        "Breakfast": {"check": "opening",  "threshold":  8 * 60, "label": "breakfast recommendation — breakfast venues should open by 8 AM"},
    }
    for venue in meal_venues:
        name = venue.get("name", "")
        hours_str = venue.get("opening_hours", "")
        section = venue.get("meal_section", "")
        cfg = configs.get(section)
        if not cfg or not hours_str:
            continue

        is_closing = cfg["check"] == "closing"
        time_re = _CLOSING_TIME_RE if is_closing else _OPENING_TIME_RE
        matches = list(time_re.finditer(hours_str))
        if not matches:
            continue

        all_mins = [
            _to_minutes(int(m.group(1)), int(m.group(2)) if m.group(2) else None, m.group(3))
            for m in matches
        ]
        best_idx = all_mins.index(max(all_mins) if is_closing else min(all_mins))
        best = all_mins[best_idx]
        if is_closing and best >= cfg["threshold"]:
            continue
        if not is_closing and best <= cfg["threshold"]:
            continue

        best_m = matches[best_idx]
        min_str = f":{best_m.group(2)}" if best_m.group(2) else ":00"
        time_str = f"{best_m.group(1)}{min_str} {best_m.group(3).upper()}"
        action = "closes" if is_closing else "opens"
        desc = (
            f"{name} {action} at {time_str} but is listed as a {cfg['label']}"
            if name else
            f"Venue {action} at {time_str} — listed as a {cfg['label']}"
        )
        section_label = _meal_section_label(section, venue)
        findings.append(Finding(
            check_id="R11",
            layer="rule",
            severity="RED",
            section=section_label,
            description=desc,
            evidence=f"{name}  Opening Hours: {hours_str}",
        ))


def _meal_section_label(meal_section: str, venue: dict) -> str:
    """Build a rich section label: 'Day 3 – Dinner Recommendations (Near Shibuya)'."""
    day = venue.get("day_number")
    area = venue.get("area", "")
    day_part  = f"Day {day} – " if day else ""
    area_part = f" ({area})" if area else ""
    return f"{day_part}{meal_section} Recommendations{area_part}"


def _r12_excessive_walking(meal_venues: list[dict], findings: list[Finding]) -> None:
    """Flag meal venues with walk times ≥ 20 min or travel times ≥ 30 min.

    Uses AI-extracted structured data so venue names are reliable regardless
    of AIG formatting conventions.
    """
    for venue in meal_venues:
        name = venue.get("name", "")
        meal_section = venue.get("meal_section", "Meal")
        section_label = _meal_section_label(meal_section, venue)

        walk = venue.get("walk_minutes")
        if walk is not None and walk >= 20:
            findings.append(Finding(
                check_id="R12",
                layer="rule",
                severity="YELLOW",
                section=section_label,
                description=(
                    f"{name}: {walk}-minute walk is too far for a meal recommendation — keep within 20 minutes on foot"
                    if name else
                    f"{walk}-minute walk listed for a meal recommendation — keep within 20 minutes on foot"
                ),
                evidence=f"{name}  Walk: {walk} minutes" if name else f"Walk: {walk} minutes",
            ))

        travel = venue.get("travel_minutes")
        if travel is not None and travel >= 30:
            findings.append(Finding(
                check_id="R12",
                layer="rule",
                severity="YELLOW",
                section=section_label,
                description=(
                    f"{name}: {travel}-minute travel time is too far for a meal recommendation — keep within 30 minutes by any mode"
                    if name else
                    f"{travel}-minute travel listed for a meal recommendation — keep within 30 minutes by any mode"
                ),
                evidence=f"{name}  Travel: {travel} minutes" if name else f"Travel: {travel} minutes",
            ))


def _r9_encoding(text: str, findings: list[Finding]) -> None:
    # � = Unicode replacement char (garbled byte during conversion)
    # - = Word/Symbol private-use glyphs (Wingdings etc.) that survive DOCX extraction
    # ￾/﻿ = reversed/regular BOM mid-document (encoding artefact)
    m = _ENC_RE.search(text)
    if m:
        start = max(0, m.start() - 30)
        snippet = text[start:m.end() + 30].strip()
        findings.append(Finding(
            check_id="R9",
            layer="rule",
            severity="YELLOW",
            section="Entire document",
            description="Broken or non-standard character encoding detected — clean up before sending to client",
            evidence=snippet or "Non-standard characters found",
        ))

def run_rule_engine(
    paragraphs: list[dict],
    maps_count: int = 0,
    meal_venues: list[dict] | None = None,
) -> list[Finding]:
    """Run all deterministic rule checks. Returns only failing checks."""
    text = paragraphs_to_text(paragraphs)
    findings: list[Finding] = []
    _r1_ai_artifacts(text, findings)
    _r2_placeholders(text, findings)
    _r3_mandatory_sections(text, findings)
    _r4_r10_day_numbers(text, findings)
    _r5_day_heading_format(text, findings)
    _r6_maps_links(text, findings, maps_count=maps_count)
    _r7_restaurant_count(text, findings)
    _r8_time_format(text, findings)
    _r9_encoding(text, findings)
    _r11_meal_timing(meal_venues or [], findings)
    _r12_excessive_walking(meal_venues or [], findings)
    return findings


# ---------------------------------------------------------------------------
# AI pass (GPT structured output)
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a professional quality control reviewer for "Bon Voyage by Marina", a luxury travel agency. \
You review All Inclusive Guides (AIGs) given to clients for real vacations. Mistakes harm clients and damage the agency's reputation.

You will receive the full text of an AIG. Review it against the checks below.
Return ONLY findings where an issue is present — do not list passing checks.
Every finding MUST include an "evidence" field: a verbatim quote from the document (max 300 chars) \
that supports the finding. If you cannot find direct supporting text, do not include the finding.

--- GLOBAL RULE: CONFIDENCE CALIBRATION ---
Many checks require you to apply external knowledge (opening days, sunset times, seasonal windows, \
reservation requirements, transport pass rules, travel times, etc.). Apply this rule to ALL such findings:
- Use severity RED only when you are HIGHLY CONFIDENT the information is wrong.
- Use severity YELLOW when you have reasonable doubt but are not certain.
- If uncertain, add the phrase "— verify before sending" to the description.
Do NOT invent facts. A false positive erodes trust in this QC process. If you cannot find \
direct evidence or are not confident in your external knowledge, do not include the finding.

--- CHECKS ---

[A1] DIETARY VIOLATIONS — severity: RED
Are any food recommendations incompatible with the dietary preferences stated in the Client Information section?
A vegetarian client CAN dine at a restaurant that also serves non-vegetarian food, as long as the venue offers vegetarian options. Only flag if: (a) the restaurant is described as exclusively non-vegetarian with no vegetarian options, or (b) a specific non-vegetarian dish is directly recommended to a vegetarian/vegan/halal client.
Other examples of genuine violations: shellfish dish recommended to a client with a shellfish allergy; pork dish recommended to a halal client; non-halal restaurant with no halal options for a halal client.
Do NOT flag a restaurant merely because it is a multi-cuisine or mixed-menu restaurant.
Evidence: Quote the dietary preference and the specific offending recommendation.

[A2] EMERGENCY CONTACTS — severity: RED
Does Safety & Emergency Contacts have real, destination-specific numbers? \
A generic "Emergency: 112" with no country context fails. "Police: 110 (Japan)" passes.
Evidence: Quote the safety section.

[A3] WRONG DESTINATION CONTENT — severity: RED
Does any day's content describe a different destination? Flag copy-paste errors from other AIGs.
Evidence: Quote the mismatched text.

[A4] GUIDE TITLE — severity: YELLOW
Is the title creative and destination-specific? "All Inclusive Guide – Japan" fails. \
"Sakura Adventures: Journey Through Japan" passes.
Evidence: Quote the title.

[A5] PACKING LIST — severity: YELLOW
Is the packing list specific to this trip's season, destination, and activities — not a generic traveller list?
Evidence: Quote sample packing items.

[A6] FULL OPENING HOURS — severity: YELLOW
Do restaurant and place entries include specific opening times (not just "Mon–Sun" with no times)?
Evidence: Quote an entry with incomplete hours.

[A7] MEAL PROXIMITY — severity: YELLOW
Are lunch places near the day's attractions, and dinner places near the hotel? \
Flag any dinner recommendation where the stated walking time is 20 minutes or more — \
that is not "near the hotel". Examples: "40-minute walk", "25 min walk", "20 min on foot". \
Also flag travel times of 30 minutes or more by any mode (taxi, transit, etc.). \
Also flag lunch venues clearly on the opposite side of the city from that day's attractions.
Evidence: Quote the proximity claim.

[A8] MUST-TRY DISHES COVERAGE — severity: YELLOW
Does the Must-Try Local Dishes section cover every destination city in the itinerary?
Evidence: Quote the section heading or note which city is absent.

[A9] GETTING AROUND COVERAGE — severity: YELLOW
Does Getting Around have transport options for each city visited?
Evidence: Quote the section or note the missing city.

[A10] CULTURAL ETIQUETTE — severity: YELLOW
Is the Cultural Etiquette section specific to the destinations, or generic boilerplate?
Evidence: Quote a sample from the section.

[A11] THANK YOU PERSONALIZATION — severity: YELLOW
Does the Thank You page use the client's actual name(s) from Client Information, or is it template text?
Evidence: Quote the Thank You text.

[A12] OVERALL COHERENCE — severity: YELLOW
Are there coherence issues: missing day narratives, broken sentences, duplicate paragraphs, or disjointed flow?
Evidence: Quote the problematic text.

[A13] HOURS LOGIC — severity: RED
Flag time ranges where the end time is BEFORE the start time within the same AM/PM period. \
E.g., "11 PM – 10 PM", "3:00 PM – 1:00 PM", "9:00 AM – 7:00 AM". \
Also flag hours strings that look malformed or self-contradictory — e.g., \
"9:00 AM – 10:30 PM | 11:30 PM – 5 PM" (the second session ends before it starts in PM). \
Note: midnight starts (12:xx AM) and identical start/end times are already caught by separate rule checks — do NOT duplicate those findings here. Only flag end-before-start in the same AM/PM period.
Check EVERY venue's hours string — do not stop at the first issue.
Evidence: Quote the exact hours string.

[A14] MEAL VENUE HOURS — severity: RED
Check meal venues across all three meal types against these thresholds:
- DINNER (any heading containing "Dinner"): closing time must be 9:00 PM or later. A dinner venue closing before 9 PM is a timing conflict.
- LUNCH (any heading containing "Lunch"): closing time must be 2:00 PM or later. A lunch venue closing at 1 PM or earlier is a timing conflict.
- BREAKFAST (any heading containing "Breakfast"): opening time must be 8:00 AM or earlier. A breakfast venue that doesn't open until 9 AM or later is a timing conflict.
Check EVERY venue under every meal section — do not stop at the first issue.
Evidence: Quote the venue name and its exact hours.

[A15] SUNSET / TIME-OF-DAY ACCURACY — severity: RED
Are sunset viewing times, golden hour visits, and similar recommendations accurate for the destination \
and travel month? E.g., suggesting a 6 PM sunset in Amsterdam in July is wrong (sunset ≈ 9:45 PM there).
Evidence: Quote the recommendation and the stated time.

[A16] TRAVEL TIME PLAUSIBILITY — severity: RED
Are stated travel times between named locations realistic given the actual geography?
Evidence: Quote the exact travel time claim.

[A17] TRANSPORT PASS ACCURACY — severity: RED
Are transport pass coverage claims accurate? E.g., Nozomi and Mizuho bullet trains historically require \
a supplement on standard JR Passes — any claim that these are "covered under JR Pass" must be flagged, \
even if qualified with "(if applicable)". The "(if applicable)" hedge is NOT sufficient — \
the client needs explicit guidance on whether their booking requires a supplement, not a vague disclaimer.
Evidence: Quote the pass coverage claim.

[A18] IMPORTANT PLACES COMPLETENESS — severity: RED
For each hotel in the itinerary, does Important Places Around Your Stay include ALL of the following?
- Grocery store (with opening hours)
- Pharmacy — preferably 24-hour; if not, note the hours explicitly so clients know when it closes
- Hospital or emergency clinic capable of handling serious medical situations (not just a GP clinic)
- Distance or travel time from the hotel for each essential service
- A Google Maps link for each essential service (missing links on hospitals/pharmacies is a safety issue)
Flag any hotel stay where one or more of these requirements is absent or insufficient.
Evidence: Quote what is present, and specify exactly what is missing.

[A19] VENUE TYPE FOR MEAL — severity: YELLOW
Are dinner/lunch recommendations appropriate for that meal? A coffee café or dessert shop listed \
as a dinner recommendation is a weak choice even if technically open at dinner time.
Evidence: Quote the venue description and the meal slot it appears under.

[A20] DISTANCE REFERENCE ANCHORING — severity: YELLOW
Are distance/time references anchored to the correct preceding attraction or hotel? \
E.g., "18 min from Attraction B" for a lunch listed before visiting Attraction B in that day's sequence.
Evidence: Quote the distance claim in context.

[A21] ACTIVITY DAY-OF-WEEK VALIDATION — severity: RED
If the guide lists opening days for an attraction, verify they include the scheduled day of the week. \
If opening days are not listed in the guide but you know with high confidence that the attraction is \
closed on that day (e.g., a well-known museum closed Mondays), flag it. \
Do not flag if you are merely guessing — apply the confidence rule above.
Evidence: Quote the day heading and the attraction's closure information.

[A22] RESERVATION DEPENDENCIES — severity: RED
Some attractions require advance booking, timed-entry tickets, or reservations. If the guide recommends \
such a venue without telling the client to book in advance, flag it. \
Flag only venues where reservations are KNOWN to be required or strongly recommended — not merely popular. \
Do NOT flag open-entry attractions, free temples, street markets, or general-access sights. \
Examples that genuinely require pre-booking: capacity-limited experiences (hot air balloon, glacier walk, \
cooking class), timed-entry monuments (Ghibli Museum, TeamLab, Colosseum), high-demand restaurants \
with no walk-in policy, cable cars or ferries with fixed departure slots. \
Do NOT flag: Eiffel Tower (walkable floors are open-access), Senso-ji Temple (free, open-access), \
busy bazaars, or any attraction where advance booking is optional rather than required.
Evidence: Quote the venue recommendation and confirm there is no booking note.

[A23] ARRIVAL / DEPARTURE DAY LOGIC — severity: RED
Arrival day: Check whether the planned activities are realistic given the stated or implied arrival time. \
If the client arrives in the afternoon or evening, a packed morning of sightseeing is wrong. \
Departure day: Check whether activities leave sufficient time to reach the airport or station. \
If a flight is at 6 PM and the last activity ends at 4 PM with a 1-hour transfer, that is dangerously tight. \
Flag any arrival or departure day where the activity schedule is incompatible with travel logistics.
Evidence: Quote the arrival/departure information and the conflicting activity schedule.

[A24] SEASONAL ACCURACY — severity: RED
Verify that seasonal experiences, natural phenomena, festivals, and weather-dependent recommendations \
align with the actual travel month stated in the itinerary. \
Examples of failures: tulip fields in Amsterdam in July (peak is April), autumn foliage in Japan in September \
(peak is late October–November), Northern Lights in Iceland in June (midnight sun, no aurora visible), \
cherry blossoms in Japan in July (peak is late March–April), whale watching in a location that is off-season. \
Use your knowledge of seasonal windows to flag recommendations that would disappoint or mislead the client.
Evidence: Quote the recommendation and state why it conflicts with the travel month.

[A26] PREMIUM FINISH — severity: YELLOW
This is a luxury travel guide going to a paying client. Flag any of the following finish issues:
- Spelling mistakes or grammatical errors in the guide text
- Broken or garbled characters (e.g., "family￾friendly", "caf?" instead of "café")
- Inconsistent capitalisation (e.g., "Day 3" in one place, "day 3" in another)
- Inconsistent emoji usage (some venues have 🍴 and ⏰, others have none, with no pattern)
- Obvious formatting problems: extra blank lines mid-paragraph, misaligned sections, section headings \
  that appear mid-sentence
- Any text that looks unpolished or rushed for a premium client document
Do NOT use A26 for: opening hours errors (AM/PM issues, midnight times, wrong durations) — those belong in R8/A13.
Evidence: Quote the specific text that has the issue.

[A27] DAY FLOW LOGIC — severity: RED
Review each day's activity sequence for practical geographic and logistical flow:
- Do activities follow a logical geographic order, or does the route backtrack unnecessarily \
  (e.g., north side of city in morning, south for lunch, north again for afternoon)?
- Are travel times between consecutive activities realistic and included where needed?
- Is there sufficient buffer time between activities, meals, and transit?
- Is the energy pacing reasonable (e.g., not 6 major sites before lunch, then nothing)?
- Can the full day realistically be completed given opening times, travel times, and meal stops?
Flag days where the sequence would leave a real client frustrated, exhausted, or unable to complete the plan.
Evidence: Quote the day heading and the specific activity sequence or timing that creates the problem.

[A25] REAL-WORLD EXECUTABILITY — severity: RED
Review each day as if you were personally taking this trip as a client. \
Flag anything that is technically possible but practically unreasonable for a real traveller: \
- Walking 40+ minutes to a dinner venue when alternatives exist nearby \
- Scheduling 6 or more major attractions in a single day with no realistic time buffer \
- Excessive backtracking across a city (visiting north, then south, then north again) \
- An activity sequence where travel time between consecutive venues makes the schedule unworkable \
- Recommending a venue that requires significant effort (long drive, permit, physical difficulty) without flagging it \
This is the check that catches itineraries that look fine on paper but would frustrate a real client.
Evidence: Quote the specific day, activity sequence, or venue that creates the problem.

[A28] DAY DATE CONTINUITY — severity: YELLOW
Check that consecutive day headings have dates that follow in sequence without unexplained gaps. \
If Day 0 is June 13 and Day 1 jumps to July 4 (21 days later) with no explanation, that will confuse \
the client — flag it. A gap is acceptable only if the guide explicitly labels it (e.g., "Educational Program", \
"Pre-Trip Arrival"). A 1-day gap between days is normal (travel day with no separate heading); \
flag gaps of 3 or more days that have no contextual explanation in the guide.
Evidence: Quote the two day headings with their dates.

[A29] IRRELEVANT CONTENT FOR TRIP MODE — severity: YELLOW
Does the guide include content that doesn't apply to how this client is actually travelling? \
Examples: parking availability and driving directions throughout an itinerary that uses only public transport; \
car rental suggestions when the client has no vehicle; cruise-specific sections in a land-only tour. \
If the client's transport mode is clear from the itinerary, flag venue entries or sections that \
assume a different mode without explanation.
Evidence: Quote the irrelevant content and specify the mismatch.

[A30] RESTAURANT DESTINATION-CULTURE FIT — severity: YELLOW
Are the restaurant recommendations culturally matched to the destination? \
If a client is visiting Japan, Italy, or Morocco for 7+ days and the majority of dinner recommendations \
are from a completely different cuisine (e.g., predominantly Indian or generic "international" options \
throughout a Japan itinerary), the guide is not giving the client an authentic local experience. \
Exception: if the client has dietary restrictions (vegetarian, vegan, halal, allergy) that genuinely \
limit local options, a higher proportion of non-local venues is acceptable — do not flag in those cases. \
Flag only when there is a clear and systematic over-reliance on non-local cuisines with no dietary justification.
Evidence: List the proportion of non-local vs local restaurant recommendations across the guide.

--- OUTPUT ---
Return JSON only — no prose before or after.
The "description" field must be a specific, actionable sentence naming the exact venue/section/value that failed and why — never a generic statement like "there is an issue here". The "evidence" field must be a verbatim copy-paste from the document, not a paraphrase.

Example of a GOOD finding (specific, names the venue and exact hours):
{
  "check_id": "A14",
  "severity": "RED",
  "section": "Day 5 – Dinner (Kawaguchiko)",
  "description": "Momijitei-Hoto closes at 5:00 PM but is listed as a dinner recommendation — clients will arrive after it has closed.",
  "evidence": "🍴 Momijitei-Hoto  ⏰ Opening Hours: 10:30 AM – 5:00 PM"
}


Full response shape:
{
  "findings": [ ...zero or more finding objects as shown above... ],
  "narratives": {
    "overall": "e.g. Strong guide with 2 RED issues that must be fixed before sending. The most critical are a restaurant timing conflict on Day 3 and missing emergency numbers.",
    "days": "e.g. Day pacing is well structured for a family trip. Day 3 has a dinner venue that closes at 5 PM.",
    "restaurants": "e.g. Vegetarian options are present throughout. Two dinner venues have hours that end before dinner service.",
    "static_sections": "e.g. Packing list is season-specific. Safety section has real country numbers. Cultural etiquette is destination-specific."
  }
}
"""


def run_ai_pass(text: str, client: AIClient | None = None) -> tuple[list[Finding], dict]:
    """Run AI verification pass. Returns (findings, narratives)."""
    if client is None:
        client = get_ai_client()

    raw = client.complete_json(text, max_tokens=8000, system=_SYSTEM_PROMPT)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(f"AI returned invalid JSON: {e}\n\nRaw output:\n{raw[:500]}") from e

    findings = []
    for item in data.get("findings") or []:
        check_id = item.get("check_id", "")
        severity = item.get("severity", "YELLOW")
        if severity not in ("RED", "YELLOW"):
            severity = "YELLOW"
        findings.append(Finding(
            check_id=check_id,
            layer="ai",
            severity=severity,
            section=item.get("section", ""),
            description=item.get("description", ""),
            evidence=item.get("evidence", ""),
        ))

    narratives = data.get("narratives") or {}
    for key in ("overall", "days", "restaurants", "static_sections"):
        narratives.setdefault(key, "")

    return findings, narratives


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def verify(docx_bytes: bytes) -> VerifyResult:
    """Run full verification pipeline on an AIG DOCX. Returns VerifyResult."""
    paragraphs = extract_paragraphs(docx_bytes)
    text = paragraphs_to_text(paragraphs)
    maps_count = _count_maps_hyperlinks(docx_bytes)

    # One shared client for both extraction passes
    client = get_ai_client()

    # Extract structured meal venue data for R11 (reliable names + hours from AI)
    meal_data = extract_from_text(client, text, VERIFY_EXTRACTION_PROMPT)
    meal_venues = meal_data.get("restaurants", [])

    rule_findings = run_rule_engine(paragraphs, maps_count=maps_count, meal_venues=meal_venues)
    ai_findings, narratives = run_ai_pass(text, client=client)

    all_findings = rule_findings + ai_findings
    red_count = sum(1 for f in all_findings if f.severity == "RED")
    yellow_count = sum(1 for f in all_findings if f.severity == "YELLOW")

    # Checks that passed = total checks - findings with issues
    all_check_ids = {
        "R1", "R2", "R3", "R4", "R5", "R6", "R7", "R8", "R9", "R10", "R11", "R12",
        "A1", "A2", "A3", "A4", "A5", "A6", "A7", "A8", "A9", "A10",
        "A11", "A12", "A13", "A14", "A15", "A16", "A17", "A18", "A19", "A20",
        "A21", "A22", "A23", "A24", "A25", "A26", "A27",
        "A28", "A29", "A30",
    }
    failed_ids = {f.check_id for f in all_findings}
    passed_count = len(all_check_ids - failed_ids)

    cost = client.cost_usd
    return VerifyResult(
        findings=all_findings,
        narratives=narratives,
        meta={
            "red_count": red_count,
            "yellow_count": yellow_count,
            "passed_count": passed_count,
            "model": client.model,
            "prompt_tokens": client.usage["prompt_tokens"],
            "completion_tokens": client.usage["completion_tokens"],
            "cost_usd": cost,
        },
    )
