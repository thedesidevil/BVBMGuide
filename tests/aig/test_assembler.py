import json
import pytest
from pathlib import Path
from docx import Document

from src.aig.assembler import assemble


def _write_fixture_days(tmp_path: Path):
    days_dir = tmp_path / "days"
    days_dir.mkdir()
    (days_dir / "day_01.json").write_text(json.dumps({
        "day_number": 1, "date": "2026-04-19", "title": "Arrival in Amsterdam",
        "overnight_city": "Amsterdam", "overnight_hotel": "Hotel V",
        "is_cruise_day": False,
        "activities": [{"name": "Check in", "vivid_description": "Cozy hotel.", "hours": None,
                        "entry_fee": None, "recommended_duration": None,
                        "source": "ai_estimated", "maps_url": "https://maps.google.com",
                        "travel_leg": {"from_location": "Airport", "mode": "taxi",
                                       "duration": "~30 min", "tip": None}}],
        "lunch": [], "dinner": [], "transport_note": None,
    }))


def _write_fixture_sections(tmp_path: Path):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    (sections_dir / "cover.json").write_text(json.dumps(
        {"title": "Golden Canals", "subtitle": "Amsterdam 2026"}
    ))
    (sections_dir / "client_info.json").write_text(json.dumps({
        "client_names": ["Ana"], "num_guests": "2 adults",
        "departure_city": "Mumbai", "dietary_preferences": None,
        "trip_summary": [{"city": "Amsterdam", "dates": "Apr 19–22",
                          "hotel": "Hotel V", "hotel_maps_link": "https://maps.google.com",
                          "transport": "Public Transport"}],
    }))
    for section in ["must_try_dishes", "souvenir_guide", "getting_around",
                    "cultural_etiquette", "mobile_connectivity", "safety_contacts",
                    "health_vaccination", "packing_list"]:
        (sections_dir / f"{section}.json").write_text(json.dumps({"cities": []}))
    (sections_dir / "thank_you.json").write_text(json.dumps(
        {"text": "Thank you for choosing Bon Voyage by Marina."}
    ))


class TestAssemble:
    def test_creates_docx_file(self, tmp_path):
        _write_fixture_days(tmp_path)
        _write_fixture_sections(tmp_path)
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)
        assert output.exists()
        assert output.stat().st_size > 0

    def test_output_is_valid_docx(self, tmp_path):
        _write_fixture_days(tmp_path)
        _write_fixture_sections(tmp_path)
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_cover_title_in_document(self, tmp_path):
        _write_fixture_days(tmp_path)
        _write_fixture_sections(tmp_path)
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)
        doc = Document(str(output))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Golden Canals" in text

    def test_day_heading_in_document(self, tmp_path):
        _write_fixture_days(tmp_path)
        _write_fixture_sections(tmp_path)
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)
        doc = Document(str(output))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Arrival in Amsterdam" in text

    def test_thank_you_page_last(self, tmp_path):
        _write_fixture_days(tmp_path)
        _write_fixture_sections(tmp_path)
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)
        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Bon Voyage by Marina" in texts[-1]

    def test_missing_section_file_does_not_crash(self, tmp_path):
        _write_fixture_days(tmp_path)
        sections_dir = tmp_path / "sections"
        sections_dir.mkdir()
        # Only cover and thank_you — everything else missing
        (sections_dir / "cover.json").write_text(json.dumps({"title": "T", "subtitle": "S"}))
        (sections_dir / "thank_you.json").write_text(json.dumps({"text": "Thanks!"}))
        output = tmp_path / "guide.docx"
        assemble(tmp_path, output)  # must not raise
        assert output.exists()
