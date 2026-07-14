"""
Golden regression tests for hotel-options document generation.

For each Excel input, the test generates a .docx with mocked external
dependencies (AI, Google Places enricher) and compares it against the
checked-in reference document.  If future code changes alter the output,
the assertion fails and the reference must be reviewed and re-approved.

Comparison strategy (see _doc_snapshot):
  - Short paragraphs (<= 200 chars): compared by exact text + style name.
  - Long paragraphs (> 200 chars): compared as "[LONG:<style>]"
    so AI-generated hotel descriptions don't break the test.
  - Table cells: compared exactly (prices, plan labels, hotel names).
  - Photo runs: have empty .text, so they are implicitly skipped.
"""
from __future__ import annotations

import base64
import io
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from src.hotel_options.models import EnrichedHotel
from src.library.ui.services.hotel_options_service import generate_doc

BASE = Path(__file__).parent

# ---------------------------------------------------------------------------
# Fake photo — minimal valid 1×1 JPEG so python-docx embeds an image
# (photo runs have empty .text and are therefore skipped by the snapshot).
# ---------------------------------------------------------------------------
_FAKE_PHOTO: bytes = base64.b64decode(
    "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsLDBkS"
    "Ew8UHRofHh0aHBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/wAALCAAB"
    "AAEBAREA/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAw"
    "IEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkK"
    "FhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4"
    "eXqDhIWGh4iJipOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU"
    "1dbX2Nna4eLj5OXm5+jp6vHy8/T19vf4+fr/2gAIAQEAAD8A+9P/2Q=="
)

# Placeholder description: >200 chars so the snapshot records [LONG:Normal].
_LONG_DESC = "X" * 201


# ---------------------------------------------------------------------------
# Helper: build an EnrichedHotel with only the fields that affect output.
# ---------------------------------------------------------------------------

def _enriched(
    official_name: str,
    rating: float,
    rating_count: int,
    cancellation: str,
    meal_type: str,
    category: str,
    room_type: str,
    dates: str = "",
    description: str = _LONG_DESC,
) -> EnrichedHotel:
    return EnrichedHotel(
        official_name=official_name,
        address="1 Test Street",
        phone="+00000000000",
        rating=rating,
        rating_count=rating_count,
        maps_url="https://maps.google.com",
        photo_bytes=_FAKE_PHOTO,
        description=description,
        cancellation=cancellation,
        meal_type=meal_type,
        category=category,
        dates=dates,
        room_type=room_type,
    )


# ---------------------------------------------------------------------------
# EnrichedHotel fixtures — keyed by the hotel.name coming out of the parser.
# Key facts sourced from the checked-in reference documents.
# ---------------------------------------------------------------------------

_LONDON_HOTELS: dict[str, EnrichedHotel] = {
    # Appears in both Wimbledon sections; enriched once from the first section.
    "Travelodge London Raynes Park": _enriched(
        official_name="Travelodge London Raynes Park",
        rating=3.9, rating_count=640,
        cancellation="Non-refundable",
        meal_type="Breakfast included",
        category="",              # no 🏨 Category line in reference
        room_type="1 x Double Room",
    ),
    "The Lodge Hotel - Putney": _enriched(
        official_name="The Lodge Hotel Putney",
        rating=4.4, rating_count=696,
        cancellation="Free cancellation till 26 Jun",
        meal_type="Breakfast included",
        category="4-Star",
        room_type="Classic Double Room",
    ),
    "Copthorne Tara Hotel London Kensington": _enriched(
        official_name="Copthorne Tara Hotel London Kensington",
        rating=3.9, rating_count=9627,
        cancellation="Free cancellation till 29 Jun",
        meal_type="Breakfast included",
        category="4-Star",
        room_type="Standard Double Room",
    ),
    # Parser strips "(Recommended)" from the name and sets recommended=True.
    "Hilton London Kensington": _enriched(
        official_name="Hilton London Kensington",
        rating=3.8, rating_count=5278,
        cancellation="Free cancellation till 29 Jun",
        meal_type="Breakfast included",
        category="4-Star",
        room_type="DOUBLE HILTON GUESTROOM",
    ),
    "Park Plaza Westminster Bridge London": _enriched(
        official_name="Park Plaza London Westminster Bridge",
        rating=4.5, rating_count=13941,
        cancellation="Free cancellation till 28 Jun",
        meal_type="Breakfast included",
        category="4-Star",
        room_type="Superior Room-1 King",
    ),
}

_JAPAN_HOTELS: dict[str, EnrichedHotel] = {
    "Via Inn Prime Kyotoeki Hachijoguchi": _enriched(
        official_name="Via Inn Prime Kyoto Station Hachijo Exit",
        rating=4.3, rating_count=1792,
        cancellation="Free cancellation till 14 Nov",
        meal_type="Breakfast included",
        category="3",
        room_type="Deluxe Double Room",
        dates="Nov 17 -22",
    ),
    # Shared across Plan A and Plan B; enriched once from Plan A.
    "HOTEL MYSTAYS Fuji Onsen Resort": _enriched(
        official_name="HOTEL MYSTAYS Fuji Onsen Resort",
        rating=4.3, rating_count=3156,
        cancellation="Free cancellation till 18 Nov",
        meal_type="Breakfast included",
        category="3",
        room_type="comfort twin room",
        dates="Nov 22 -24",
    ),
    # Shared across Plan A and Plan B; enriched once from Plan A.
    "Grand Nikko Tokyo Daiba": _enriched(
        official_name="GRAND NIKKO TOKYO DAIBA",
        rating=4.4, rating_count=9327,
        cancellation="Free cancellation till 21 Nov",
        meal_type="Breakfast included",
        category="5",
        room_type="Premier Floor Superior Double, Non Smoking",
        dates="Nov 24 -26",
    ),
    "Mercure Kyoto Station": _enriched(
        official_name="Mercure Kyoto Station",
        rating=4.4, rating_count=1785,
        cancellation="",           # no 🔄 Cancellation line in reference
        meal_type="Breakfast included",
        category="4",
        room_type="Superior Room with 1 King bed",
        dates="Nov 17 -22",
    ),
}


# ---------------------------------------------------------------------------
# Mock factories
# ---------------------------------------------------------------------------

def _make_ai_client(*responses: str) -> MagicMock:
    """Return a mock AI client whose .complete() yields responses in order."""
    mock = MagicMock()
    it = iter(responses)
    mock.complete.side_effect = lambda prompt, **kw: next(it)
    mock.cost_usd = 0.0
    return mock


def _make_enricher_mock(hotel_map: dict[str, EnrichedHotel]) -> MagicMock:
    """
    Mock the _enricher module.

    - check_hotels_exist: returns a place_id for every hotel that is in
      hotel_map; returns None for unknown hotels (they are then skipped
      by the generator and do not appear in the output).
    - enrich_hotel / enrich_hotel_multi_segment: look up by hotel.name.
    - fetch_destination_photo: always returns _FAKE_PHOTO.
    """
    mock = MagicMock()

    def _check(names, dest, key):
        return {n: (f"pid_{i}" if n in hotel_map else None)
                for i, n in enumerate(names)}

    mock.check_hotels_exist.side_effect = _check
    mock.enrich_hotel.side_effect = (
        lambda hotel, place_id, dest, key, ai: hotel_map[hotel.name]
    )
    mock.enrich_hotel_multi_segment.side_effect = (
        lambda hotels, place_id, dest, key, ai: hotel_map[hotels[0].name]
    )
    mock.fetch_destination_photo.return_value = _FAKE_PHOTO
    mock.place_id_from_maps_url.return_value = None
    return mock


def _make_storage_mock() -> MagicMock:
    """StorageBackend mock: read_json returns None → CodeStore uses seed codes only."""
    mock = MagicMock()
    mock.read_json.return_value = None
    return mock


# ---------------------------------------------------------------------------
# Snapshot helper
# ---------------------------------------------------------------------------

def _doc_snapshot(docx_bytes: bytes) -> dict:
    """
    Extract a comparable representation of the document:

    paragraphs:
      - Empty paragraphs are skipped (photos and spacers produce no text).
      - Paragraphs ≤ 200 chars: "[<style>] <text>" (exact comparison).
      - Paragraphs > 200 chars: "[LONG:<style>]"   (presence-only, content ignored).
    tables:
      - Every non-empty row as a list of stripped cell texts (exact comparison).
    """
    doc = Document(io.BytesIO(docx_bytes))

    paragraphs: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if len(text) > 200:
            paragraphs.append(f"[LONG:{p.style.name}]")
        else:
            paragraphs.append(f"[{p.style.name}] {text}")

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        t_data: list[list[str]] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                t_data.append(cells)
        if t_data:
            tables.append(t_data)

    return {"paragraphs": paragraphs, "tables": tables}


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

_SERVICE = "src.library.ui.services.hotel_options_service"


def test_london_golden():
    """
    London (no-plans / grouped-by-sections layout).

    AI call order in generate_doc when grouped_by_sections=True:
      1. _normalize_section_labels  → cleaned section header list
      2. _infer_destination_from_labels → "London"
      3. _format_stay_requirements  → stay requirements string
    """
    xlsx = (
        BASE / "resources/input/DO NOT SHARE_ Bushan_no plans.xlsx"
    ).read_bytes()
    ref = (
        BASE / "resources/output/Hotel Options - London.docx"
    ).read_bytes()

    enricher_mock = _make_enricher_mock(_LONDON_HOTELS)
    ai_mock = _make_ai_client(
        '["Wimbledon (Jun 28 - Jul 4)", "Wimbledon (Jun 28 - Jul 1)", "Central London (Jul 1 - Jul 4)"]',
        "London",
        "Refundable Booking • Breakfast Included",
    )
    storage_mock = _make_storage_mock()

    with (
        patch(f"{_SERVICE}.get_ai_client", return_value=ai_mock),
        patch(f"{_SERVICE}._enricher", enricher_mock),
    ):
        result_bytes, _, _ = generate_doc(
            xlsx,
            "DO NOT SHARE_ Bushan_no plans.xlsx",
            {},
            {},
            storage_mock,
            "fake_key",
        )

    assert _doc_snapshot(result_bytes) == _doc_snapshot(ref)


def test_japan_golden():
    """
    Japan (explicit Plan A / Plan B layout).

    AI call order in generate_doc when grouped_by_sections=False:
      1. _format_stay_requirements → stay requirements string
      (normalize_section_labels and _infer_destination are skipped)
    """
    xlsx = (
        BASE / "resources/input"
        / "DO NOT SHARE- Rochak_Accommodation Options_Japan_v2_plans.xlsx"
    ).read_bytes()
    ref = (
        BASE / "resources/output/Hotel Options - Japan.docx"
    ).read_bytes()

    enricher_mock = _make_enricher_mock(_JAPAN_HOTELS)
    ai_mock = _make_ai_client("Refundable Booking • Breakfast Included")
    storage_mock = _make_storage_mock()

    with (
        patch(f"{_SERVICE}.get_ai_client", return_value=ai_mock),
        patch(f"{_SERVICE}._enricher", enricher_mock),
    ):
        result_bytes, _, _ = generate_doc(
            xlsx,
            "DO NOT SHARE- Rochak_Accommodation Options_Japan_v2_plans.xlsx",
            {},
            {},
            storage_mock,
            "fake_key",
        )

    assert _doc_snapshot(result_bytes) == _doc_snapshot(ref)
