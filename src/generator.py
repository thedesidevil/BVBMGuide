"""
AIG Generator — written from scratch based on docs/ knowledge base.

Generates All Inclusive Guides section by section using a TDD-style QC loop:
each section is validated immediately after generation and retried up to 2 times
before being accepted. Unresolved failures are appended to a QC Issues page.

Missing input fields are rendered as [PLACEHOLDER] strings with red highlighting
so a human reviewer can spot and fix them quickly.
"""

from __future__ import annotations

import json
import re
from copy import deepcopy
from datetime import datetime, timedelta
from pathlib import Path
from typing import Callable, Optional, TypeVar

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from rich.console import Console

from .ai_provider import AIClient
from .maps import maps_url
from .models import (
    AIGDay, AIGDocument, ClientInfoPage, ClientInfoRow,
    FoodPreferences, ImportantPlace, ImportantPlacesSection,
    ItineraryDay, MealRecommendation, QCIssue, Restaurant, UserInputs,
    PLACEHOLDER_CLIENT_NAME, PLACEHOLDER_HOTEL,
    PLACEHOLDER_NUM_PAX, PLACEHOLDER_DEPARTURE, PLACEHOLDER_DATES,
)
from .qc import DocumentQC, SectionQC

console = Console()

T = TypeVar("T")

BVBM_HEADER = (
    "BON VOYAGE BY MARINA\n"
    "Tower E 2503, Panchshil Towers, Wagholi, Pune, M.H. 412207\n"
    "Shop Act MH: 104297592503   Contact Number: +91-86000-15316\n"
    "Udyam: UDYAM-MH-26-0889139   Email ID: hello@bonvoyagebymarina.com"
)

MORNING_KEYWORDS = {
    "sunrise", "morning", "breakfast", "market", "early", "dawn", "brunch",
    "check-in", "arrive", "arrival", "transfer", "flight", "depart",
}
EVENING_KEYWORDS = {
    "sunset", "evening", "dinner", "night", "nightlife", "rooftop", "bar",
    "show", "performance", "farewell", "dusk",
}


class AIGGenerator:
    """Generates a complete AIG from user inputs and a pre-ranked restaurant list."""

    def __init__(
        self,
        user_inputs: UserInputs,
        restaurants: list[Restaurant],
        client: AIClient,
    ):
        self.inputs = user_inputs
        self.restaurants = restaurants
        self.client = client
        self.section_qc = SectionQC(ai_client=client)
        self._unresolved: list[QCIssue] = []
        # Tracks restaurant names used per day: index = day_number-1
        self._restaurant_history: list[list[str]] = []

    # ------------------------------------------------------------------
    # Public entry points
    # ------------------------------------------------------------------

    def generate(self) -> AIGDocument:
        dest = self.inputs.itinerary.destination
        hotel = self.inputs.hotel_name or PLACEHOLDER_HOTEL
        client_names = self.inputs.client_names or [PLACEHOLDER_CLIENT_NAME]
        trip_dates = self.inputs.trip_dates or PLACEHOLDER_DATES

        console.print("\n[bold blue]Generating AIG...[/bold blue]")

        # Title
        console.print("  [dim]→ Title[/dim]")
        title = self._generate_with_qc(
            "Cover Page Title",
            lambda: self._gen_title(dest, trip_dates, client_names),
            lambda t: self.section_qc.check_title(t, dest),
            lambda t, issues: self._gen_title(dest, trip_dates, client_names, issues),
        )

        # Client info page
        console.print("  [dim]→ Client Information Page[/dim]")
        client_info = self._generate_with_qc(
            "Client Information Page",
            lambda: self._gen_client_info(client_names, trip_dates, hotel, dest),
            lambda c: self.section_qc.check_client_info(c),
            lambda c, issues: self._gen_client_info(client_names, trip_dates, hotel, dest, issues),
        )

        # Day-wise itinerary
        days: list[AIGDay] = []
        start_dt = self._parse_start_date(trip_dates)
        for iday in self.inputs.itinerary.days:
            date_str, day_name = self._date_for_day(iday.day_number, start_dt)
            console.print(f"  [dim]→ Day {iday.day_number}[/dim]")
            day = self._generate_with_qc(
                f"Day {iday.day_number}",
                lambda d=iday, ds=date_str, dn=day_name: self._gen_day(d, ds, dn),
                lambda day_obj: (
                    self.section_qc.check_day_header(day_obj)
                    + self.section_qc.check_day_restaurants(
                        day_obj, self._recent_restaurants(day_obj.day_number)
                    )
                ),
                lambda day_obj, issues, d=iday, ds=date_str, dn=day_name: self._gen_day(
                    d, ds, dn, issues
                ),
            )
            days.append(day)
            self._record_restaurants(day)

        # Supplementary sections
        console.print("  [dim]→ Important Places[/dim]")
        important_places = self._gen_important_places_all_cities(dest, hotel)

        console.print("  [dim]→ Must-Try Dishes[/dim]")
        must_try = self._gen_supplementary(
            "must_try_dishes",
            self._prompt_must_try(dest, self.inputs.food_preferences),
        )

        console.print("  [dim]→ Getting Around[/dim]")
        getting_around = self._gen_supplementary(
            "getting_around",
            self._prompt_getting_around(dest, self.inputs),
        )

        console.print("  [dim]→ Souvenir Guide[/dim]")
        souvenir = self._gen_supplementary(
            "souvenir_guide",
            self._prompt_souvenir(dest),
        )

        console.print("  [dim]→ Cultural Etiquette[/dim]")
        etiquette = self._gen_supplementary(
            "cultural_etiquette",
            self._prompt_etiquette(dest),
        )

        console.print("  [dim]→ Packing List[/dim]")
        packing = self._gen_supplementary(
            "packing_list",
            self._prompt_packing(dest, trip_dates, self.inputs),
        )

        console.print("  [dim]→ Mobile Connectivity[/dim]")
        mobile = self._gen_supplementary(
            "mobile_connectivity",
            self._prompt_mobile(dest),
        )

        console.print("  [dim]→ Safety & Emergency[/dim]")
        safety = self._gen_supplementary(
            "safety_contacts",
            self._prompt_safety(dest),
        )

        console.print("  [dim]→ Health & Vaccination[/dim]")
        health = self._gen_supplementary(
            "health_vaccination",
            self._prompt_health(dest),
        )

        console.print("  [dim]→ Thank You Page[/dim]")
        thank_you = self._gen_thank_you(client_names, dest)

        doc = AIGDocument(
            title=title,
            subtitle=f"A Personalised Travel Guide by Bon Voyage by Marina",
            client_names=client_names,
            destination=dest,
            hotel_name=hotel,
            hotel_maps_link=maps_url(hotel, dest),
            trip_dates=trip_dates,
            client_info=client_info,
            days=days,
            must_try_dishes=must_try,
            important_places_structured=important_places,
            souvenir_guide=souvenir,
            cultural_etiquette=etiquette,
            packing_list=packing,
            safety_contacts=safety,
            mobile_connectivity=mobile,
            health_vaccination=health,
            getting_around=getting_around,
            thank_you_page=thank_you,
            unresolved_qc_issues=self._unresolved,
        )

        # Final document-level QC (informational only)
        final_issues = DocumentQC().check(doc)
        if final_issues:
            console.print("\n[yellow]Final QC issues:[/yellow]")
            for issue in final_issues:
                console.print(f"  [yellow]⚠[/yellow] {issue.section}: {', '.join(issue.descriptions)}")
        else:
            console.print("\n[green]✓ All mandatory QC checks passed[/green]")

        return doc

    # ------------------------------------------------------------------
    # TDD loop
    # ------------------------------------------------------------------

    def _generate_with_qc(
        self,
        label: str,
        generate_fn: Callable[[], T],
        validate_fn: Callable[[T], list[str]],
        fix_fn: Optional[Callable[[T, list[str]], T]] = None,
    ) -> T:
        result = generate_fn()
        for attempt in range(2):
            issues = validate_fn(result)
            if not issues:
                return result
            console.print(
                f"    [yellow]QC retry {attempt + 1}/2 for {label}:[/yellow] "
                + "; ".join(issues[:2])
            )
            if fix_fn:
                result = fix_fn(result, issues)
            else:
                break
        remaining = validate_fn(result)
        if remaining:
            self._unresolved.append(QCIssue(section=label, descriptions=remaining))
            console.print(f"    [red]Unresolved QC issues in {label} — added to QC Issues page[/red]")
        return result

    # ------------------------------------------------------------------
    # Section generators
    # ------------------------------------------------------------------

    def _gen_title(
        self,
        destination: str,
        trip_dates: str,
        client_names: list[str],
        prior_issues: Optional[list[str]] = None,
    ) -> str:
        issues_note = ""
        if prior_issues:
            issues_note = f"\n\nPrevious attempt failed QC:\n" + "\n".join(f"- {i}" for i in prior_issues)

        prompt = f"""You are naming a premium travel guide for Bon Voyage by Marina.

Destination: {destination}
Travel dates: {trip_dates}
Client names: {", ".join(client_names)}

Create a SINGLE creative, evocative title for their All Inclusive Guide.
Rules:
- 3 to 8 words
- Destination-specific and evocative (e.g. "Sakura Trails: Japan Through Your Eyes", "La Dolce Vita: An Italian Journey")
- NOT generic (never "All Inclusive Guide" or "Travel Guide")
- May use a colon to separate a poetic phrase from a location reference
- Premium, warm, aspirational tone{issues_note}

Return ONLY the title text, nothing else."""

        return self._ai_complete(prompt).strip().strip('"').strip("'")

    def _gen_client_info(
        self,
        client_names: list[str],
        trip_dates: str,
        hotel: str,
        destination: str,
        prior_issues: Optional[list[str]] = None,
    ) -> ClientInfoPage:
        # Determine cities from itinerary days
        itinerary = self.inputs.itinerary
        cities = self._extract_cities(destination)

        rows: list[ClientInfoRow] = []
        start_dt = self._parse_start_date(trip_dates)

        # Group days by city if possible; fallback to one row per hotel
        # For simplicity, create one row per unique city derived from the itinerary
        if len(cities) <= 1:
            cities = [destination]

        days_per_city = max(1, len(itinerary.days) // max(len(cities), 1))
        for i, city in enumerate(cities):
            day_start = i * days_per_city + 1
            day_end = (i + 1) * days_per_city if i < len(cities) - 1 else len(itinerary.days)
            if start_dt:
                d1 = start_dt + timedelta(days=day_start - 1)
                d2 = start_dt + timedelta(days=day_end - 1)
                city_dates = f"{d1.strftime('%d %b')} – {d2.strftime('%d %b %Y')}"
            else:
                city_dates = PLACEHOLDER_DATES

            city_hotel = hotel if i == 0 else PLACEHOLDER_HOTEL
            rows.append(ClientInfoRow(
                city=city,
                dates=city_dates,
                hotel=city_hotel,
                hotel_maps_link=maps_url(city_hotel, city),
                transport="Taxi" if i == 0 else "Train / Taxi",
            ))

        prefs = self.inputs.food_preferences
        dietary = ", ".join(prefs.dietary_restrictions) if prefs.dietary_restrictions else "No restrictions"

        num_pax = PLACEHOLDER_NUM_PAX
        departure = self.inputs.additional_notes or PLACEHOLDER_DEPARTURE

        return ClientInfoPage(
            client_names=client_names if client_names else [PLACEHOLDER_CLIENT_NAME],
            num_pax=num_pax,
            departure_city=departure,
            dietary_preferences=dietary,
            trip_summary_rows=rows,
        )

    def _gen_day(
        self,
        iday: ItineraryDay,
        date_str: str,
        day_name: str,
        prior_issues: Optional[list[str]] = None,
    ) -> AIGDay:
        dest = self.inputs.itinerary.destination
        hotel = self.inputs.hotel_name or PLACEHOLDER_HOTEL
        prefs = self.inputs.food_preferences
        issues_note = ""
        if prior_issues:
            issues_note = "\n\nFix these issues from the previous attempt:\n" + "\n".join(
                f"- {i}" for i in prior_issues
            )

        # Classify activities by time of day
        morning, afternoon, evening = [], [], []
        for act in iday.activities:
            slot = self._classify_time(act.name)
            desc = self._gen_activity_description(act.name, dest, iday.day_number == 1)
            entry = f"📍 {act.name}\n{desc}"
            if slot == "morning":
                morning.append(entry)
            elif slot == "evening":
                evening.append(entry)
            else:
                afternoon.append(entry)

        # Balance: if all fell into one bucket, spread them
        all_acts = morning + afternoon + evening
        if len(all_acts) >= 2 and not morning:
            morning.append(all_acts[0])
            afternoon = all_acts[1:]
            evening = []
        if len(all_acts) >= 3 and not evening:
            evening = [afternoon.pop()] if afternoon else []

        # Special Day 1: prepend arrival logistics
        if iday.day_number == 1:
            arrival = self._gen_arrival_note(dest, hotel)
            morning.insert(0, arrival)

        # Restaurants
        lunch = self._assign_restaurants(iday.day_number, "lunch", prefs)
        dinner = self._assign_restaurants(iday.day_number, "dinner", prefs)

        overnight = self._extract_city_for_day(iday.day_number) or dest

        return AIGDay(
            day_number=iday.day_number,
            date=date_str,
            day_name=day_name,
            title=iday.title,
            morning_activities=morning,
            afternoon_activities=afternoon,
            evening_activities=evening,
            lunch_recommendations=lunch,
            dinner_recommendations=dinner,
            overnight_location=overnight,
        )

    def _assign_restaurants(
        self, day_number: int, meal_type: str, prefs: FoodPreferences
    ) -> MealRecommendation:
        dest = self.inputs.itinerary.destination
        hotel = self.inputs.hotel_name or ""
        dietary = set(r.lower() for r in prefs.dietary_restrictions)
        allergies = set(a.lower() for a in prefs.allergies)
        recent = self._recent_restaurants(day_number)

        def is_compatible(r: Restaurant) -> bool:
            if dietary and "vegetarian" in dietary and not r.vegetarian_friendly:
                return False
            for allergy in allergies:
                if allergy in " ".join(r.cuisine_type).lower():
                    return False
            return True

        def rotation_ok(r: Restaurant) -> bool:
            return recent.count(r.name) < 2

        compatible = [r for r in self.restaurants if is_compatible(r) and rotation_ok(r)]
        selected = compatible[:3]

        # Curate with AI if library can't supply 3
        while len(selected) < 3:
            curated = self._curate_restaurant(dest, hotel, meal_type, prefs, selected)
            selected.append(curated)

        return MealRecommendation(meal_type=meal_type, options=selected[:3])

    def _curate_restaurant(
        self,
        destination: str,
        hotel: str,
        meal_type: str,
        prefs: FoodPreferences,
        existing: list[Restaurant],
    ) -> Restaurant:
        dietary = ", ".join(prefs.dietary_restrictions) or "none"
        existing_names = [r.name for r in existing]

        prompt = f"""Suggest ONE real restaurant in {destination} for {meal_type}.

Requirements:
- Must be a real, operating restaurant with Google rating ≥4.3
- ≤20 minutes from hotel "{hotel}"
- Suitable for dietary restrictions: {dietary}
- NOT one of: {existing_names}
- For {meal_type}: {"near attractions/central area" if meal_type == "lunch" else "near the hotel"}

Return JSON with these exact fields:
{{
  "name": "Restaurant Name",
  "cuisine_type": ["Cuisine"],
  "ambience": "One sentence describing atmosphere",
  "must_try_dishes": ["Dish 1", "Dish 2"],
  "hours": "HH:MM AM – HH:MM PM (closed Mondays)",
  "price_range": "₹800–1,200 per person",
  "distance_from_hotel": "12 min taxi from {hotel}"
}}

Return ONLY valid JSON."""

        try:
            response = self._ai_complete(prompt)
            text = response.strip()
            if text.startswith("```"):
                text = re.sub(r"^```(?:json)?\n?", "", text)
                text = re.sub(r"\n?```$", "", text)
            data = json.loads(text)
            name = data.get("name", "Local Restaurant")
            return Restaurant(
                name=name,
                location=destination,
                cuisine_type=data.get("cuisine_type", []),
                ambience=data.get("ambience", ""),
                must_try_dishes=data.get("must_try_dishes", []),
                hours=data.get("hours", ""),
                price_range=data.get("price_range", ""),
                distance_from_hotel=data.get("distance_from_hotel", f"≤20 min from {hotel}"),
                google_maps_link=maps_url(name, destination),
                vegetarian_friendly="vegetarian" in dietary.lower(),
                is_curated=True,
            )
        except Exception:
            return Restaurant(
                name="[LOCAL RESTAURANT]",
                location=destination,
                google_maps_link=maps_url("restaurant", destination),
                is_curated=True,
                distance_from_hotel=f"Near {hotel}",
            )

    def _gen_important_places_all_cities(
        self, destination: str, hotel: str
    ) -> list[ImportantPlacesSection]:
        cities = self._extract_cities(destination)
        if not cities:
            cities = [destination]

        result: list[ImportantPlacesSection] = []
        for city in cities:
            city_hotel = hotel if city == destination else PLACEHOLDER_HOTEL
            section = self._generate_with_qc(
                f"Important Places ({city})",
                lambda c=city, h=city_hotel: self._gen_important_places_city(c, h),
                lambda s: self.section_qc.check_important_places(s),
                lambda s, issues, c=city, h=city_hotel: self._gen_important_places_city(c, h, issues),
            )
            result.append(section)
        return result

    def _gen_important_places_city(
        self,
        city: str,
        hotel: str,
        prior_issues: Optional[list[str]] = None,
    ) -> ImportantPlacesSection:
        issues_note = ""
        if prior_issues:
            issues_note = "\n\nFix these issues:\n" + "\n".join(f"- {i}" for i in prior_issues)

        prompt = f"""List essential nearby places for a tourist staying at "{hotel}" in {city}.

Provide REAL places. Return JSON:
{{
  "places": [
    {{
      "name": "Place Name",
      "category": "grocery",
      "address": "Street address",
      "hours": "7:00 AM – 10:00 PM",
      "distance_from_hotel": "8 min walk",
      "phone": "+XX XXXX XXXX"
    }}
  ]
}}

Required categories (include ≥1 of each):
- "grocery": supermarket or convenience store, within 10 min walk
- "pharmacy": pharmacy with extended hours, within 10 min walk
- "hospital": 24/7 emergency hospital, reasonable distance{issues_note}

Return ONLY valid JSON."""

        try:
            response = self._ai_complete(prompt)
            text = response.strip()
            if text.startswith("```"):
                text = re.sub(r"^```(?:json)?\n?", "", text)
                text = re.sub(r"\n?```$", "", text)
            data = json.loads(text)
            places = []
            for p in data.get("places", []):
                name = p.get("name", "")
                places.append(ImportantPlace(
                    name=name,
                    category=p.get("category", ""),
                    address=p.get("address", ""),
                    hours=p.get("hours", ""),
                    distance_from_hotel=p.get("distance_from_hotel", ""),
                    google_maps_link=maps_url(name, city),
                    phone=p.get("phone", ""),
                ))
        except Exception:
            places = []

        return ImportantPlacesSection(city=city, hotel_name=hotel, places=places)

    def _gen_supplementary(self, section_name: str, prompt: str) -> str:
        prefs = self.inputs.food_preferences
        dest = self.inputs.itinerary.destination

        return self._generate_with_qc(
            section_name,
            lambda: self._ai_complete(prompt),
            lambda content: self.section_qc.check_supplementary(
                section_name, content, dest, prefs
            ),
            lambda content, issues: self._ai_complete(
                prompt
                + f"\n\nPrevious attempt had these issues — fix them:\n"
                + "\n".join(f"- {i}" for i in issues)
            ),
        )

    def _gen_thank_you(self, client_names: list[str], destination: str) -> str:
        names = ", ".join(client_names) if client_names else "Traveller"
        prompt = f"""Write a warm, personal Thank You page for a Bon Voyage by Marina travel guide.

Client names: {names}
Destination: {destination}

The message should:
- Thank the client for choosing Bon Voyage by Marina
- Express genuine warmth and personal connection
- Wish them an amazing journey with specific destination reference
- Invite them to reach out if they need anything during the trip
- End with the BVBM team signature

Keep it to 3–4 short paragraphs. Warm, premium, personal — not corporate."""
        return self._ai_complete(prompt)

    # ------------------------------------------------------------------
    # AI section prompts
    # ------------------------------------------------------------------

    def _gen_activity_description(self, activity_name: str, destination: str, is_day1: bool) -> str:
        prompt = f"""Write 2-3 vivid, immersive sentences describing "{activity_name}" in {destination}.

Style:
- Friendly, storytelling tone like a knowledgeable local guide
- Include sensory details: light, atmosphere, what makes it special
- Practical: mention what to do/see/experience
- Avoid generic filler

Return ONLY the description text, no labels."""
        return self._ai_complete(prompt, max_tokens=200).strip()

    def _gen_arrival_note(self, destination: str, hotel: str) -> str:
        return (
            f"✈️ Arrival in {destination}\n"
            f"● Upon landing, proceed through immigration and collect your baggage.\n"
            f"● 🚖 Take a pre-booked taxi or arrange one at the airport to reach {hotel}.\n"
            f"● Check in, freshen up, and get ready to explore {destination}!"
        )

    def _prompt_must_try(self, destination: str, prefs: FoodPreferences) -> str:
        dietary = ", ".join(prefs.dietary_restrictions) or "none"
        return f"""Write the "🍽️ Must-Try Local Dishes" section for a travel guide to {destination}.

Client dietary restrictions: {dietary}

Format:
- Group dishes by category: Appetizers / Street Food, Main Dishes, Desserts & Drinks
- For each dish: name in bold, 1-2 sentence description, where to try it
- Mark vegetarian options with ○
- Include 8-12 dishes total
- If dietary restrictions apply, flag incompatible dishes and provide alternatives
- Authentic to {destination}, not generic "international" food

Write in a warm, appetising tone that makes the reader excited to eat."""

    def _prompt_getting_around(self, destination: str, inputs: UserInputs) -> str:
        return f"""Write the "🚆 Getting Around {destination}" section for a travel guide.

Trip context: {inputs.trip_dates or "multi-day trip"}
Hotel: {inputs.hotel_name or destination}

Cover:
- Primary transport modes used in this trip (taxi, train, metro, tuk-tuk, etc.)
- How to book / hail each transport type
- Estimated costs in INR
- Recommended apps or booking tools
- Any passes or cards that save money (e.g. day passes, tourist cards)
- Tips to avoid scams or overcharging
- Specific advice for getting from the airport to the hotel

Be practical and specific to {destination}. Not generic travel advice."""

    def _prompt_souvenir(self, destination: str) -> str:
        return f"""Write the "🎁 Souvenir Shopping Guide" for a travel guide to {destination}.

Cover:
- 5-8 specific items to buy that are authentic to {destination}
- WHERE to buy each (name specific markets, shops, or districts with Maps links in format: https://www.google.com/maps/search/?api=1&query=PLACE_NAME+{destination.replace(" ", "+")})
- Tips on authenticity (how to spot fakes)
- Haggling tips where applicable
- Payment methods accepted
- Best time to visit shops

Be location-specific. Avoid generic souvenir advice."""

    def _prompt_etiquette(self, destination: str) -> str:
        return f"""Write the "🌍 Cultural Etiquette & Useful Local Phrases" section for {destination}.

Format exactly as:

**Dos**
● [point]
● [point]
● [point]
● [point]

**Don'ts**
● [point]
● [point]
● [point]
● [point]

**Essential Phrases**
Greetings:
● [English] → [Local Language] → [Pronunciation if needed]

Politeness:
● [English] → [Local Language]

Requests & Questions:
● [English] → [Local Language]

Dietary Needs:
● "I am vegetarian" → [Local Language]
● "I have a nut allergy" → [Local Language]

Emergency:
● "Help!" → [Local Language]
● "Call the police" → [Local Language]

Be accurate and relevant to {destination}. Only include phrases actually used there."""

    def _prompt_packing(self, destination: str, trip_dates: str, inputs: UserInputs) -> str:
        dietary = ", ".join(inputs.food_preferences.dietary_restrictions) or "none"
        return f"""Write the "🎒 Tailored Packing List" for a trip to {destination}.

Trip dates: {trip_dates}
Dietary: {dietary}

Organise by:
👚 Essentials (universal travel items)
🌤️ {destination}-specific (weather, terrain, activities)
🎭 Cultural / Religious requirements
🧴 Personal & Health items

For each item: one line, specific not generic (e.g. "reef-safe sunscreen SPF50" not just "sunscreen").
Include any destination-specific must-haves (adapter type, specific medications, etc.).
Keep it practical — 20-30 items total."""

    def _prompt_mobile(self, destination: str) -> str:
        return f"""Write the "📱 Mobile Connectivity & SIM Card Guide" for {destination}.

Cover:
- Best eSIM options (Airalo, Saily, Nomad) with estimated data costs
- Best local SIM card providers with network quality notes
- Where to buy (airport vs. city shops)
- Documents required
- Coverage quality across the trip destinations
- Recommended data plans for a 2-week trip

Format clearly with provider names in bold. Be accurate for {destination}."""

    def _prompt_safety(self, destination: str) -> str:
        return f"""Write the "🚨 Personal Safety & Emergency Contacts" section for {destination}.

Include:
1. 5-6 practical safety tips numbered, with sub-bullets for detail
2. Emergency numbers (police, ambulance, fire — with country code)
3. Indian Embassy / High Commission in {destination}: address, phone, email
4. Nearest major hospital to the hotel area
5. 3-4 common tourist scams to watch out for, with how to avoid them

Be specific to {destination}. Not generic safety advice."""

    def _prompt_health(self, destination: str) -> str:
        return f"""Write the "💉 Health & Vaccination Guidance" for travel to {destination}.

Cover:
- Routine vaccines to ensure are up to date
- Destination-specific recommended vaccines (only if genuinely relevant — don't over-recommend)
- Any required vaccines (e.g. yellow fever for certain regions)
- Food and water safety tips specific to {destination}
- Common health risks and how to avoid them
- Recommended medications to carry (anti-diarrhoeal, antihistamine, etc.)
- A short, warm closing line

Tone: reassuring and practical, not alarmist. Clearly distinguish required vs. recommended."""

    # ------------------------------------------------------------------
    # Date utilities
    # ------------------------------------------------------------------

    def _parse_start_date(self, trip_dates: str) -> Optional[datetime]:
        if not trip_dates or "[" in trip_dates:
            return None
        # Try patterns: "Apr 19 - May 5, 2025", "19 Apr 2025", "2025-04-19", etc.
        patterns = [
            r"(\d{1,2}\s+\w+\s+\d{4})",           # "19 Apr 2025"
            r"(\w+\s+\d{1,2},?\s*\d{4})",          # "April 19, 2025"
            r"(\w+\s+\d{1,2})\s*[-–]",             # "Apr 19 -" (year assumed)
            r"(\d{4}-\d{2}-\d{2})",                 # ISO format
        ]
        year_hint = re.search(r"\b(20\d{2})\b", trip_dates)
        year = int(year_hint.group(1)) if year_hint else datetime.now().year

        for pattern in patterns:
            m = re.search(pattern, trip_dates)
            if m:
                date_str = m.group(1).strip().rstrip(",")
                for fmt in ("%d %b %Y", "%B %d %Y", "%b %d %Y", "%Y-%m-%d", "%b %d", "%B %d"):
                    try:
                        dt = datetime.strptime(date_str, fmt)
                        if dt.year == 1900:
                            dt = dt.replace(year=year)
                        return dt
                    except ValueError:
                        continue
        return None

    def _date_for_day(self, day_number: int, start_dt: Optional[datetime]) -> tuple[str, str]:
        if start_dt:
            dt = start_dt + timedelta(days=day_number - 1)
            return dt.strftime("%d %b %Y"), dt.strftime("%A")
        return PLACEHOLDER_DATES, "[DAY NAME]"

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _classify_time(self, activity_name: str) -> str:
        lower = activity_name.lower()
        if any(kw in lower for kw in MORNING_KEYWORDS):
            return "morning"
        if any(kw in lower for kw in EVENING_KEYWORDS):
            return "evening"
        return "afternoon"

    def _extract_cities(self, destination: str) -> list[str]:
        """Try to derive multiple cities from the itinerary days' titles."""
        cities = []
        seen: set[str] = set()
        for day in self.inputs.itinerary.days:
            # Look for "in CityName" patterns in day titles
            m = re.search(r"\bin\s+([A-Z][a-zA-Z\s]+?)(?:\s*[|,–\-]|$)", day.title)
            if m:
                city = m.group(1).strip()
                if city not in seen:
                    seen.add(city)
                    cities.append(city)
        return cities if cities else [destination]

    def _extract_city_for_day(self, day_number: int) -> Optional[str]:
        days = self.inputs.itinerary.days
        if day_number <= len(days):
            m = re.search(r"\bin\s+([A-Z][a-zA-Z\s]+?)(?:\s*[|,–\-]|$)", days[day_number - 1].title)
            if m:
                return m.group(1).strip()
        return None

    def _recent_restaurants(self, day_number: int) -> list[str]:
        """Return restaurant names used in the 3 days before day_number."""
        recent: list[str] = []
        for i in range(max(0, day_number - 4), day_number - 1):
            if i < len(self._restaurant_history):
                recent.extend(self._restaurant_history[i])
        return recent

    def _record_restaurants(self, day: AIGDay) -> None:
        names = (
            [r.name for r in day.lunch_recommendations.options]
            + [r.name for r in day.dinner_recommendations.options]
        )
        self._restaurant_history.append(names)

    def _ai_complete(self, prompt: str, max_tokens: int = 1500) -> str:
        return self.client.complete(prompt)

    # ------------------------------------------------------------------
    # DOCX rendering
    # ------------------------------------------------------------------

    def save_docx(self, doc: AIGDocument, output_path: Path) -> None:
        output_path = Path(output_path)
        word = Document()

        # Page margins
        for section in word.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Default font
        style = word.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # BVBM header on every page
        self._add_page_header(word)

        # Cover page
        self._add_cover_page(word, doc)

        # TOC
        word.add_page_break()
        self._add_toc(word, doc)

        # Client info page
        word.add_page_break()
        self._add_client_info_page(word, doc)

        # Day sections
        for day in doc.days:
            word.add_page_break()
            self._add_day_section(word, day, doc.destination)

        # Supplementary sections (in QC checklist order)
        supp_sections = [
            ("🏥 Important Places Around Your Stay", None, doc.important_places_structured),
            ("🎁 Souvenir Shopping Guide", doc.souvenir_guide, None),
            ("🍽️ Must-Try Local Dishes", doc.must_try_dishes, None),
            ("🚆 Getting Around", doc.getting_around, None),
            ("🌍 Cultural Etiquette & Local Phrases", doc.cultural_etiquette, None),
            ("🎒 Tailored Packing List", doc.packing_list, None),
            ("📱 Mobile Connectivity Guide", doc.mobile_connectivity, None),
            ("🚨 Safety & Emergency Contacts", doc.safety_contacts, None),
            ("💉 Health & Vaccination Guidance", doc.health_vaccination, None),
        ]

        for title, text_content, structured_content in supp_sections:
            word.add_page_break()
            self._add_section_header(word, title)
            if structured_content is not None:
                self._add_important_places(word, structured_content)
            elif text_content:
                self._add_text_section(word, text_content)
            else:
                p = word.add_paragraph()
                self._add_placeholder_run(p, f"[{title.upper()} — CONTENT MISSING]")

        # Thank you page
        word.add_page_break()
        self._add_section_header(word, "🙏 Thank You")
        if doc.thank_you_page:
            self._add_text_section(word, doc.thank_you_page)

        # QC Issues page (only if there are unresolved issues)
        if doc.unresolved_qc_issues:
            word.add_page_break()
            self._add_qc_issues_page(word, doc.unresolved_qc_issues)

        word.save(str(output_path))
        console.print(f"\n[bold green]✓ Saved:[/bold green] {output_path}")

    # ------------------------------------------------------------------
    # DOCX helpers
    # ------------------------------------------------------------------

    def _add_page_header(self, word: Document) -> None:
        for section in word.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()
            run = p.add_run(BVBM_HEADER)
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _add_cover_page(self, word: Document, doc: AIGDocument) -> None:
        word.add_paragraph()  # spacing
        word.add_paragraph()

        title_p = word.add_paragraph()
        title_p.alignment = 1  # center
        run = title_p.add_run(doc.title)
        run.font.name = "Arial"
        run.font.size = Pt(28)
        run.font.bold = True
        self._check_and_highlight(title_p)

        if doc.subtitle:
            sub_p = word.add_paragraph()
            sub_p.alignment = 1
            sub_run = sub_p.add_run(doc.subtitle)
            sub_run.font.size = Pt(14)
            sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        word.add_paragraph()

        dest_p = word.add_paragraph()
        dest_p.alignment = 1
        dest_run = dest_p.add_run(f"✈️  {doc.destination}")
        dest_run.font.size = Pt(16)
        dest_run.font.bold = True

        if doc.trip_dates:
            dates_p = word.add_paragraph()
            dates_p.alignment = 1
            dates_run = dates_p.add_run(doc.trip_dates)
            dates_run.font.size = Pt(13)
            self._check_and_highlight(dates_p)

        if doc.client_names:
            names_p = word.add_paragraph()
            names_p.alignment = 1
            names_run = names_p.add_run(f"Prepared for: {', '.join(doc.client_names)}")
            names_run.font.size = Pt(13)
            self._check_and_highlight(names_p)

    def _add_toc(self, word: Document, doc: AIGDocument) -> None:
        h = word.add_heading("Table of Contents", level=1)
        h.runs[0].font.size = Pt(18)
        word.add_paragraph()

        # Note for user
        note = word.add_paragraph()
        note_run = note.add_run("(Right-click and select 'Update Field' in Word to refresh page numbers)")
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        word.add_paragraph()

        entries = []
        for day in doc.days:
            day_name = day.day_name or ""
            date = day.date or ""
            entries.append(f"📅  Day {day.day_number}: {day_name}, {date} – {day.title}")

        entries += [
            "🏥  Important Places Around Your Stay",
            "🎁  Souvenir Shopping Guide",
            "🍽️  Must-Try Local Dishes",
            "🚆  Getting Around",
            "🌍  Cultural Etiquette & Local Phrases",
            "🎒  Tailored Packing List",
            "📱  Mobile Connectivity Guide",
            "🚨  Safety & Emergency Contacts",
            "💉  Health & Vaccination Guidance",
            "🙏  Thank You",
        ]

        for entry in entries:
            p = word.add_paragraph(entry, style="List Paragraph")
            p.paragraph_format.space_after = Pt(4)

    def _add_client_info_page(self, word: Document, doc: AIGDocument) -> None:
        self._add_section_header(word, "👤 Client Information")
        word.add_paragraph()

        ci = doc.client_info
        if not ci:
            p = word.add_paragraph()
            self._add_placeholder_run(p, PLACEHOLDER_CLIENT_NAME)
            return

        def add_info_row(label: str, value: Optional[str]) -> None:
            p = word.add_paragraph()
            label_run = p.add_run(f"{label}: ")
            label_run.font.bold = True
            if value:
                val_run = p.add_run(value)
                if "[" in value:
                    val_run.font.highlight_color = WD_COLOR_INDEX.RED
                    val_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                self._add_placeholder_run(p, f"[{label.upper()}]")

        add_info_row("Client Name(s)", ", ".join(ci.client_names))
        add_info_row("Number of Guests", ci.num_pax)
        add_info_row("Departure City", ci.departure_city)
        add_info_row("Dietary Preferences", ci.dietary_preferences)
        word.add_paragraph()

        # Trip summary table
        if ci.trip_summary_rows:
            h = word.add_paragraph("Trip Summary")
            h.runs[0].font.bold = True
            h.runs[0].font.size = Pt(13)
            word.add_paragraph()

            table = word.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["City", "Dates", "Hotel", "Transport", "Maps"]
            for i, hdr in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = hdr
                cell.paragraphs[0].runs[0].font.bold = True

            for row in ci.trip_summary_rows:
                cells = table.add_row().cells
                cells[0].text = row.city
                cells[1].text = row.dates
                hotel_p = cells[2].paragraphs[0]
                hotel_run = hotel_p.add_run(row.hotel)
                if "[" in row.hotel:
                    hotel_run.font.highlight_color = WD_COLOR_INDEX.RED
                    hotel_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                cells[3].text = row.transport
                maps_p = cells[4].paragraphs[0]
                self._add_hyperlink(maps_p, "📍 View Map", row.hotel_maps_link)

    def _add_day_section(self, word: Document, day: AIGDay, destination: str) -> None:
        # Day heading
        date_str = day.date or PLACEHOLDER_DATES
        day_name = day.day_name or "[DAY]"
        heading_text = f"Day {day.day_number}: {day_name}, {date_str} – {day.title}"
        h = word.add_heading(heading_text, level=1)
        h.runs[0].font.size = Pt(16)
        self._check_and_highlight(h)
        word.add_paragraph()

        # Time blocks
        for block_label, activities in [
            ("🌅 Morning", day.morning_activities),
            ("☀️ Afternoon", day.afternoon_activities),
            ("🌙 Evening", day.evening_activities),
        ]:
            if not activities:
                continue
            block_h = word.add_paragraph(block_label)
            block_h.runs[0].font.bold = True
            block_h.runs[0].font.size = Pt(13)
            for activity_text in activities:
                for line in activity_text.split("\n"):
                    p = word.add_paragraph(line, style="List Bullet")
                    p.paragraph_format.left_indent = Inches(0.3)
                    self._check_and_highlight(p)
            word.add_paragraph()

        # Meals
        for meal in (day.lunch_recommendations, day.dinner_recommendations):
            meal_label = f"🍽️ {meal.meal_type.capitalize()} Recommendations"
            meal_h = word.add_paragraph(meal_label)
            meal_h.runs[0].font.bold = True
            meal_h.runs[0].font.size = Pt(13)
            word.add_paragraph()
            for restaurant in meal.options:
                self._add_restaurant_entry(word, restaurant, destination)
            word.add_paragraph()

        # Overnight
        if day.overnight_location:
            overnight_p = word.add_paragraph()
            overnight_run = overnight_p.add_run(f"✅ Stay Overnight in {day.overnight_location}")
            overnight_run.font.bold = True

    def _add_restaurant_entry(self, word: Document, r: Restaurant, destination: str) -> None:
        # Restaurant name as hyperlink
        name_p = word.add_paragraph()
        name_p.paragraph_format.space_before = Pt(6)
        url = r.google_maps_link or maps_url(r.name, destination)
        self._add_hyperlink(name_p, r.name, url, bold=True, size=12)
        if r.is_curated:
            name_p.add_run(" ✨").font.size = Pt(10)

        if r.ambience:
            p = word.add_paragraph(f"● {r.ambience}", style="List Bullet")
            self._check_and_highlight(p)

        if r.must_try_dishes:
            dishes = ", ".join(r.must_try_dishes)
            word.add_paragraph(f"● 🍴 Must-Try: {dishes}", style="List Bullet")

        details: list[str] = []
        if r.hours:
            details.append(f"⏰ {r.hours}")
        if r.price_range:
            details.append(f"💰 {r.price_range}")
        if r.distance_from_hotel:
            details.append(f"🚶 {r.distance_from_hotel}")
        if details:
            p = word.add_paragraph("● " + "  |  ".join(details), style="List Bullet")
            self._check_and_highlight(p)

        if r.best_for:
            word.add_paragraph(f"● Best for: {', '.join(r.best_for)}", style="List Bullet")

        word.add_paragraph()

    def _add_important_places(
        self, word: Document, sections: list[ImportantPlacesSection]
    ) -> None:
        for section in sections:
            city_h = word.add_paragraph(f"📍 {section.city} — {section.hotel_name}")
            city_h.runs[0].font.bold = True
            city_h.runs[0].font.size = Pt(13)
            self._check_and_highlight(city_h)
            word.add_paragraph()

            for category in ("grocery", "pharmacy", "hospital"):
                cat_places = [p for p in section.places if p.category == category]
                if not cat_places:
                    p = word.add_paragraph()
                    self._add_placeholder_run(p, f"[{category.upper()} — NOT FOUND]")
                    continue

                cat_labels = {
                    "grocery": "🛒 Grocery / Convenience",
                    "pharmacy": "💊 Pharmacy",
                    "hospital": "🏥 Hospital / Emergency",
                }
                cat_h = word.add_paragraph(cat_labels.get(category, category.title()))
                cat_h.runs[0].font.bold = True

                for place in cat_places:
                    name_p = word.add_paragraph()
                    url = place.google_maps_link or maps_url(place.name, section.city)
                    self._add_hyperlink(name_p, place.name, url, bold=True)
                    self._check_and_highlight(name_p)

                    details: list[str] = []
                    if place.hours:
                        details.append(f"⏰ {place.hours}")
                    if place.distance_from_hotel:
                        details.append(f"🚶 {place.distance_from_hotel}")
                    if place.phone:
                        details.append(f"📞 {place.phone}")
                    if details:
                        word.add_paragraph("● " + "  |  ".join(details), style="List Bullet")
                    if place.address:
                        word.add_paragraph(f"● 📍 {place.address}", style="List Bullet")
                    word.add_paragraph()

    def _add_text_section(self, word: Document, content: str) -> None:
        for line in content.split("\n"):
            line = line.rstrip()
            if not line:
                word.add_paragraph()
                continue
            if line.startswith("**") and line.endswith("**"):
                p = word.add_paragraph(line.strip("*"))
                if p.runs:
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
            elif line.startswith("●") or line.startswith("•"):
                p = word.add_paragraph(line, style="List Bullet")
                self._check_and_highlight(p)
            elif line.startswith("○"):
                p = word.add_paragraph(line, style="List Bullet 2")
                self._check_and_highlight(p)
            else:
                p = word.add_paragraph(line)
                self._check_and_highlight(p)

    def _add_section_header(self, word: Document, title: str) -> None:
        h = word.add_heading(title, level=1)
        h.runs[0].font.size = Pt(18)
        word.add_paragraph()

    def _add_qc_issues_page(self, word: Document, issues: list[QCIssue]) -> None:
        h = word.add_heading("⚠️ QC Issues — Human Review Required", level=1)
        h.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        h.runs[0].font.size = Pt(18)

        intro = word.add_paragraph(
            "The following issues were not resolved automatically during generation. "
            "Please review and fix each item before sending this guide to the client."
        )
        intro.runs[0].font.italic = True
        word.add_paragraph()

        for issue in issues:
            section_p = word.add_paragraph()
            section_run = section_p.add_run(f"📌 {issue.section}")
            section_run.font.bold = True
            section_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            for desc in issue.descriptions:
                p = word.add_paragraph(f"  • {desc}", style="List Bullet")
                p.runs[0].font.color.rgb = RGBColor(0x88, 0x00, 0x00)
            word.add_paragraph()

    def _add_hyperlink(
        self,
        paragraph,
        text: str,
        url: str,
        bold: bool = False,
        size: int = 11,
    ) -> None:
        """Inject a hyperlink into a paragraph using python-docx XML."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        if bold:
            b_elem = OxmlElement("w:b")
            rPr.append(b_elem)

        sz_elem = OxmlElement("w:sz")
        sz_elem.set(qn("w:val"), str(size * 2))
        rPr.append(sz_elem)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_placeholder_run(self, paragraph, text: str) -> None:
        run = paragraph.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.RED
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.bold = True

    def _check_and_highlight(self, paragraph) -> None:
        """Highlight any [PLACEHOLDER] tokens in paragraph runs red."""
        for run in paragraph.runs:
            if re.search(r"\[.+?\]", run.text):
                run.font.highlight_color = WD_COLOR_INDEX.RED
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
