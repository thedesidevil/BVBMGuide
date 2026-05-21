"""DOCX assembler for AIG documents.

Reads all stored JSONs from input/days/ and input/sections/, applies AIG
named styles, and writes the final DOCX in mandatory section order.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.aig.styles import ensure_styles

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _load(path: Path) -> dict:
    """Load a JSON file, returning {} on missing file or parse error."""
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError, UnicodeDecodeError) as exc:
        log.warning("Failed to load %s: %s", path, exc)
        return {}


def _is_empty(data: dict) -> bool:
    """Return True if data is {} or {"cities": []} (nothing to render)."""
    if not data:
        return True
    if list(data.keys()) == ["cities"] and not data["cities"]:
        return True
    return False


def _add(doc: Document, text: str, style: str) -> None:
    """Add a paragraph with the given AIG style."""
    doc.add_paragraph(text, style=style)


def _add_hr(doc: Document) -> None:
    """Add a thin horizontal rule (paragraph with bottom border)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")        # 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet_labeled(doc: Document, label: str, value: str, level: int = 0) -> None:
    """Add a bullet where label is bold and value is normal: 'label: value'."""
    style = "List Bullet 2" if level == 0 else "List Bullet 3"
    p = doc.add_paragraph(style=style)
    r1 = p.add_run(label + ": ")
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r2 = p.add_run(value)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    r2.font.bold = False
    r2.font.color.rgb = RGBColor(0, 0, 0)


def _add_bullet(doc: Document, text: str, level: int = 0, bold: bool = False) -> None:
    """Add an Arial 11pt black bullet paragraph.

    level=0 → indented bullet (List Bullet 2)
    level=1 → double-indented sub-bullet (List Bullet 3)
    """
    style = "List Bullet 2" if level == 0 else "List Bullet 3"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _format_date(date_str: str) -> str:
    """Format YYYY-MM-DD as 'Sat, Apr 19'. Returns original string if unparseable."""
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%a, %b %-d")
    except (ValueError, TypeError):
        return date_str or ""


def _format_date_long(date_str: str) -> str:
    """Format YYYY-MM-DD as 'Sun, April 19' (weekday + full month name). Returns original if unparseable."""
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%a, %B %-d")
    except (ValueError, TypeError):
        return date_str or ""


def _remove_initial_empty_paragraphs(doc: Document) -> None:
    """Remove the default empty paragraph python-docx adds to new documents."""
    for p in list(doc.paragraphs):
        if p.text.strip() == "":
            p._element.getparent().remove(p._element)
        else:
            break


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def _render_cover(doc: Document, data: dict) -> None:
    title = data.get("title", "")
    subtitle = data.get("subtitle", "")
    if title:
        _add(doc, title, "AIG Title")
    if subtitle:
        _add(doc, subtitle, "AIG Subtitle")


def _render_client_info(doc: Document, data: dict) -> None:
    _add(doc, "Client Information", "AIG Section Heading")

    names = data.get("client_names") or []
    if names:
        _add(doc, "Travellers: " + ", ".join(names), "AIG Body")

    num_guests = data.get("num_guests")
    if num_guests:
        _add(doc, f"Guests: {num_guests}", "AIG Body")

    departure = data.get("departure_city")
    if departure:
        _add(doc, f"Departing from: {departure}", "AIG Body")

    dietary = data.get("dietary_preferences")
    if dietary:
        _add(doc, f"Dietary preferences: {dietary}", "AIG Body")

    trip_summary = data.get("trip_summary") or []
    if trip_summary:
        _add(doc, "Trip Summary", "AIG Subsection")
        for row in trip_summary:
            city = row.get("city", "")
            raw_dates = row.get("dates", "")
            dates = " – ".join(_format_date(p.strip()) for p in raw_dates.split(" – ")) if raw_dates else ""
            hotel = row.get("hotel", "")
            hotel_link = row.get("hotel_maps_link", "")
            transport = row.get("transport", "")
            line_parts = [f"{city} | {dates}"]
            if hotel:
                hotel_text = f"Hotel: {hotel}"
                if hotel_link:
                    hotel_text += f" ({hotel_link})"
                line_parts.append(hotel_text)
            if transport:
                line_parts.append(f"Transport: {transport}")
            _add(doc, " | ".join(line_parts), "AIG Detail")


# ---------------------------------------------------------------------------
# Day renderers
# ---------------------------------------------------------------------------

def _render_flight_details(doc: Document, flight: dict) -> None:
    airline = flight.get("airline", "")
    if airline:
        _add(doc, "Flight Details:", "AIG Body Bold")
        _add_bullet_labeled(doc, "Airline", airline, level=0)

    for sector in (flight.get("sectors") or []):
        label = sector.get("label", "Sector")
        from_t = sector.get("from_terminal", "")
        to_t = sector.get("to_terminal", "")
        dep = sector.get("departure_time", "")
        arr = sector.get("arrival_time", "")
        dur = sector.get("duration")

        _add_bullet(doc, label + ":", level=0, bold=True)
        if from_t:
            _add_bullet(doc, f"📍 {from_t} →", level=1)
        if to_t:
            _add_bullet(doc, f"📍 {to_t}", level=1)
        if dep:
            _add_bullet(doc, f"Departure: {dep}", level=1)
        if arr:
            _add_bullet(doc, f"Arrival: {arr}", level=1)
        if dur:
            _add_bullet(doc, f"Duration: {dur}", level=1)


def _render_transfer_instructions(doc: Document, transfer: dict) -> None:
    intro = transfer.get("intro", "")
    if intro:
        _add(doc, intro, "AIG Body")

    tip = transfer.get("headline_tip", "")
    if tip:
        _add(doc, tip, "AIG Body Bold")

    route_steps = transfer.get("route_steps") or []
    if route_steps:
        _add(doc, "Route:", "AIG Body Bold")
        for step in route_steps:
            _add_bullet(doc, step, level=0)

    time_info = transfer.get("total_travel_time", "")
    cost = transfer.get("cost_inr", "")
    if time_info or cost:
        parts = []
        if time_info:
            parts.append(f"⏱️ Total travel time: {time_info}")
        if cost:
            parts.append(f"💰 {cost}")
        _add(doc, "  ".join(parts), "AIG Body")

    note = transfer.get("settle_in_note", "")
    if note:
        _add(doc, note, "AIG Body")


def _render_activity_section(doc: Document, section: dict) -> None:
    title = section.get("section_title", "")
    if title:
        _add(doc, title, "AIG Day Section")

    desc = section.get("description", "")
    if desc:
        _add(doc, desc, "AIG Body")

    place = section.get("place_name")
    maps_url_val = section.get("maps_url", "")
    if place:
        place_text = f"📍{place}"
        if maps_url_val:
            place_text += f"  ({maps_url_val})"
        _add(doc, place_text, "AIG Body")

    travel_info = section.get("travel_info")
    if travel_info:
        _add_bullet(doc, travel_info, level=0)

    hours = section.get("hours")
    if hours:
        _add_bullet(doc, f"⏰ {hours}", level=0)

    cost = section.get("cost")
    if cost:
        _add_bullet(doc, f"💰 {cost}", level=0)


def _render_return_instructions(doc: Document, return_inst: dict) -> None:
    instructions = return_inst.get("instructions", "")
    if instructions:
        _add(doc, instructions, "AIG Body")

    travel_note = return_inst.get("travel_note", "")
    if travel_note:
        _add(doc, travel_note, "AIG Body")


def _render_restaurant(doc: Document, r: dict) -> None:
    name = r.get("name", "")
    if name:
        maps_url_val = r.get("maps_url", "")
        name_line = f"🍴 {name}"
        if maps_url_val:
            name_line += f"  ({maps_url_val})"
        _add(doc, name_line, "AIG Body")

    ambience = r.get("ambience")
    if ambience:
        _add_bullet(doc, ambience, level=0)

    must_try = r.get("must_try_dishes") or r.get("must_try") or []
    if must_try:
        if isinstance(must_try, list):
            _add_bullet(doc, "Must try: " + ", ".join(str(d) for d in must_try), level=0)
        else:
            _add_bullet(doc, f"Must try: {must_try}", level=0)

    hours = r.get("hours")
    if hours:
        _add_bullet(doc, f"⏰ {hours}", level=0)

    price_inr = r.get("price_inr")
    if price_inr:
        _add_bullet(doc, f"💰 {price_inr}", level=0)

    travel_note = r.get("travel_note")
    if travel_note:
        _add_bullet(doc, f"⏱️ {travel_note}", level=0)


def _render_day(doc: Document, data: dict) -> None:
    day_number = data.get("day_number", "?")
    date = data.get("date", "")
    title = data.get("title", "")

    # Day heading: "Day 1: April 19- Arrival in Amsterdam"
    date_str = _format_date_long(date) if date else ""
    if date_str and title:
        heading = f"Day {day_number}: {date_str}- {title}"
    elif date_str:
        heading = f"Day {day_number}: {date_str}"
    elif title:
        heading = f"Day {day_number}: {title}"
    else:
        heading = f"Day {day_number}"

    _add(doc, heading, "AIG Day Heading")
    _add_hr(doc)

    is_cruise = data.get("is_cruise_day", False)
    if is_cruise:
        doc.add_paragraph()
        _add(doc, "At Sea — enjoy the cruise amenities and relax on board.", "AIG Body")
    else:
        overnight_city = data.get("overnight_city", "")

        # Flight section (arrival days)
        flight = data.get("flight_details")
        if flight:
            doc.add_paragraph()
            _add(doc, f"✈️ Arrival & Journey to {overnight_city}", "AIG Day Section")
            _render_flight_details(doc, flight)

        # Transfer section
        transfer = data.get("transfer_instructions")
        if transfer:
            _add_hr(doc)
            _add(doc, "🚆 Transfer to Hotel", "AIG Day Section")
            _render_transfer_instructions(doc, transfer)

        # Activity sections
        for section in (data.get("sections") or []):
            _add_hr(doc)
            _render_activity_section(doc, section)

        # Dinner section
        dinner = data.get("dinner") or []
        if dinner:
            loc = data.get("dinner_location", "")
            section_title = f"🍽️ Dinner Recommendations{' Near ' + loc.title() if loc else ''}"
            _add_hr(doc)
            _add(doc, section_title, "AIG Day Section")
            for r in dinner:
                doc.add_paragraph()
                _render_restaurant(doc, r)

        # Return to hotel section
        return_inst = data.get("return_instructions")
        if return_inst:
            _add_hr(doc)
            _add(doc, "🌙 Return to Hotel", "AIG Day Section")
            _render_return_instructions(doc, return_inst)

    # Overnight
    overnight_city = data.get("overnight_city", "")
    overnight_hotel = data.get("overnight_hotel", "")
    overnight_label = overnight_city or overnight_hotel
    if overnight_label:
        doc.add_paragraph()
        _add(doc, f"✅ Overnight in {overnight_label}", "AIG Overnight")


def _render_list_section(doc: Document, heading: str, data: dict) -> None:
    """Render a 'cities' list section: heading + per-city subsections with bullet items."""
    _add(doc, heading, "AIG Section Heading")
    cities = data.get("cities") or []
    for city_block in cities:
        city_name = city_block.get("city", "")
        if city_name:
            _add(doc, city_name, "AIG Subsection")
        items = city_block.get("items") or city_block.get("dishes") or []
        for item in items:
            if isinstance(item, dict):
                label = item.get("name") or item.get("dish") or str(item)
                detail = item.get("description") or item.get("detail") or ""
                _add(doc, f"• {label}", "AIG Bullet")
                if detail:
                    _add(doc, detail, "AIG Detail")
            else:
                _add(doc, f"• {item}", "AIG Bullet")


def _render_packing_list(doc: Document, data: dict) -> None:
    """Render packing list: heading + bullet items (no city grouping)."""
    _add(doc, "Tailored Packing List", "AIG Section Heading")
    items = data.get("items") or []
    if not items:
        cities = data.get("cities") or []
        for city_block in cities:
            city_name = city_block.get("city", "")
            if city_name:
                _add(doc, city_name, "AIG Subsection")
            city_items = city_block.get("items") or []
            for item in city_items:
                if isinstance(item, dict):
                    label = item.get("name") or str(item)
                    _add(doc, f"• {label}", "AIG Bullet")
                else:
                    _add(doc, f"• {item}", "AIG Bullet")
    else:
        for item in items:
            if isinstance(item, dict):
                label = item.get("name") or str(item)
                _add(doc, f"• {label}", "AIG Bullet")
            else:
                _add(doc, f"• {item}", "AIG Bullet")


def _render_important_places(doc: Document, data: dict) -> None:
    _add(doc, "Important Places Around Your Stay", "AIG Section Heading")
    cities = data.get("cities") or []
    for city_block in cities:
        city_name = city_block.get("city", "")
        if city_name:
            _add(doc, city_name, "AIG Subsection")
        places = city_block.get("places") or city_block.get("items") or []
        for place in places:
            if isinstance(place, dict):
                name = place.get("name", "")
                desc = place.get("description") or place.get("detail") or ""
                maps_url = place.get("maps_url") or ""
                if name:
                    _add(doc, f"• {name}", "AIG Bullet")
                if desc:
                    _add(doc, desc, "AIG Detail")
                if maps_url:
                    _add(doc, f"Map: {maps_url}", "AIG Detail")
            else:
                _add(doc, f"• {place}", "AIG Bullet")


def _render_thank_you(doc: Document, data: dict) -> None:
    _add(doc, "Thank You", "AIG Section Heading")
    text = data.get("text", "")
    if text:
        _add(doc, text, "AIG Body")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def assemble(input_dir: Path, output_path: Path) -> None:
    """Assemble the final AIG DOCX from stored JSON files.

    Args:
        input_dir: Directory containing days/ and sections/ subdirectories.
        output_path: Destination path for the output .docx file.
    """
    input_dir = Path(input_dir)
    days_dir = input_dir / "days"
    sections_dir = input_dir / "sections"

    doc = Document()
    ensure_styles(doc)
    _remove_initial_empty_paragraphs(doc)

    # ------------------------------------------------------------------
    # 1. Cover
    # ------------------------------------------------------------------
    cover = _load(sections_dir / "cover.json")
    if not _is_empty(cover):
        _render_cover(doc, cover)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 2. Client Information
    # ------------------------------------------------------------------
    client_info = _load(sections_dir / "client_info.json")
    if not _is_empty(client_info):
        _render_client_info(doc, client_info)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 3. Day-wise Itinerary
    # ------------------------------------------------------------------
    if days_dir.exists():
        day_files = sorted(days_dir.glob("day_*.json"))
        for day_file in day_files:
            day_data = _load(day_file)
            if day_data:
                _render_day(doc, day_data)
                doc.add_page_break()

    # ------------------------------------------------------------------
    # 4. Important Places Around Your Stay
    # ------------------------------------------------------------------
    important_places = _load(sections_dir / "important_places.json")
    if not _is_empty(important_places):
        _render_important_places(doc, important_places)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 5. Souvenir Shopping Guide
    # ------------------------------------------------------------------
    souvenir = _load(sections_dir / "souvenir_guide.json")
    if not _is_empty(souvenir):
        _render_list_section(doc, "Souvenir Shopping Guide", souvenir)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 6. Must-Try Local Dishes
    # ------------------------------------------------------------------
    must_try = _load(sections_dir / "must_try_dishes.json")
    if not _is_empty(must_try):
        _render_list_section(doc, "Must-Try Local Dishes", must_try)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 7. Getting Around
    # ------------------------------------------------------------------
    getting_around = _load(sections_dir / "getting_around.json")
    if not _is_empty(getting_around):
        _render_list_section(doc, "Getting Around", getting_around)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 8. Cultural Etiquette & Local Phrases
    # ------------------------------------------------------------------
    etiquette = _load(sections_dir / "cultural_etiquette.json")
    if not _is_empty(etiquette):
        _render_list_section(doc, "Cultural Etiquette & Local Phrases", etiquette)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 9. Packing List
    # ------------------------------------------------------------------
    packing = _load(sections_dir / "packing_list.json")
    if not _is_empty(packing):
        _render_packing_list(doc, packing)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 10. Mobile Connectivity
    # ------------------------------------------------------------------
    connectivity = _load(sections_dir / "mobile_connectivity.json")
    if not _is_empty(connectivity):
        _render_list_section(doc, "Mobile Connectivity Guide", connectivity)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 11. Safety & Emergency Contacts
    # ------------------------------------------------------------------
    safety = _load(sections_dir / "safety_contacts.json")
    if not _is_empty(safety):
        _render_list_section(doc, "Safety & Emergency Contacts", safety)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 12. Health & Vaccination Guidance
    # ------------------------------------------------------------------
    health = _load(sections_dir / "health_vaccination.json")
    if not _is_empty(health):
        _render_list_section(doc, "Health & Vaccination Guidance", health)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # 13. Thank You (always last — NO page break after)
    # ------------------------------------------------------------------
    thank_you = _load(sections_dir / "thank_you.json")
    if not _is_empty(thank_you):
        _render_thank_you(doc, thank_you)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Assembled DOCX saved to %s", output_path)
